
import os
import base64
import gc
import re
import time
import uuid
import tempfile
import jwt
import requests
from io import BytesIO
from typing import Dict, List, Tuple

import fitz
import streamlit as st
import pdf_summary
import planning_rules
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

DEFAULT_STATE = {
    "report": None,
    "sections": None,
    "word_file": None,
    "pdf_file": None,
    "last_filename": None,
    "last_error": None,
    "report_id": None,
    "active_module": "Building Regulations Review",
    "saved_projects": [],
    "report_library": [],
    "app_theme": "Dark",
    "brand_logo_bytes": None,
    "ai_confidence": None,
    "rule_engine_summary": None,
    "credit_balance": 10,
    "credit_transactions": [],
    "unlocked_reports": {},
}
for key, value in DEFAULT_STATE.items():
    if key not in st.session_state:
        st.session_state[key] = value

MAX_FILE_SIZE_MB = 20
MAX_PAGE_COUNT = 30
STARTER_MONTHLY_REVIEW_LIMIT = 10

# -----------------------------------------------------------------------------
# CREDIT / TOKEN SYSTEM
# User-facing name: Credits. Internal name can still be treated as tokens.
# Session-based foundation first; move to Supabase/Stripe for live persistence.
# -----------------------------------------------------------------------------
CREDIT_PACKS = {
    "10 Credits": {"credits": 10, "price": "£19"},
    "30 Credits": {"credits": 30, "price": "£49"},
    "75 Credits": {"credits": 75, "price": "£99"},
}

EXPORT_CREDIT_COSTS = {
    "Planning Review": {"pdf": 3, "word": 1, "planning_statement": 3, "design_access_statement": 4},
    "Building Regulations Review": {"pdf": 5, "word": 1},
}

FREE_PREVIEW_NOTE = "Analysis preview is available first. Credits are used when exports/downloads are unlocked."

ARCHLENS_API_URL = os.getenv("ARCHLENS_API_URL", "https://archlens-api.onrender.com").rstrip("/")
ARCHLENS_WEBHOOK_SECRET = os.getenv("ARCHLENS_WEBHOOK_SECRET", "archlens_secure_2026_SYDS_92838")
ARCHLENS_BUY_CREDITS_URL = os.getenv("ARCHLENS_BUY_CREDITS_URL", "https://www.sydesignstudio.co.uk/category/archlens-ai-credits")


def normalise_user_email(email: str) -> str:
    return str(email or "").strip().lower()


def api_get_credit_balance(email: str):
    clean_email = normalise_user_email(email)
    if not clean_email:
        return None
    try:
        response = requests.get(
            f"{ARCHLENS_API_URL}/user/{clean_email}",
            timeout=10,
        )
        if response.status_code == 200:
            data = response.json()
            return int(data.get("credits", 0) or 0)
    except Exception as exc:
        print("Credit balance API error:", exc)
    return None


def api_deduct_credits(email: str, amount: int, report_id: str = "", export_type: str = ""):
    clean_email = normalise_user_email(email)
    amount = int(amount or 0)
    if not clean_email:
        return {
            "success": False,
            "message": "User email not found. Please launch ArchLens from your Wix member account.",
        }
    if amount <= 0:
        return {"success": True, "credits": api_get_credit_balance(clean_email), "message": "No credits required."}

    try:
        response = requests.post(
            f"{ARCHLENS_API_URL}/deduct-credits",
            headers={
                "Content-Type": "application/json",
                "x-archlens-secret": ARCHLENS_WEBHOOK_SECRET,
            },
            json={
                "email": clean_email,
                "credits": amount,
                "reportId": report_id,
                "exportType": export_type,
                "source": "archlens_download",
            },
            timeout=15,
        )
        try:
            data = response.json()
        except Exception:
            data = {"detail": response.text}

        if response.status_code == 200 and data.get("success", True):
            return {
                "success": True,
                "credits": int(data.get("credits", data.get("new_balance", 0)) or 0),
                "message": data.get("message", f"{amount} credits used."),
            }

        return {
            "success": False,
            "message": data.get("detail") or data.get("message") or "Credit deduction failed.",
        }

    except Exception as exc:
        return {
            "success": False,
            "message": f"Could not connect to credit API: {exc}",
        }


def sync_credit_balance_from_api(email: str):
    api_balance = api_get_credit_balance(email)
    if api_balance is not None:
        st.session_state["credit_balance"] = api_balance
    return st.session_state.get("credit_balance", 0)


BUILDING_REQUIRED_HEADINGS = [
    "PROJECT CLASSIFICATION",
    "PROJECT DETAILS",
    "TOP SUMMARY",
    "DRAWING-PACK INCONSISTENCIES",
    "EXECUTIVE SUMMARY",
    "DRAWING PACK SUMMARY",
    "COMPLIANCE STATUS BY APPROVED DOCUMENT",
    "KEY RISKS",
    "MISSING INFORMATION",
    "RECOMMENDED ACTIONS",
    "BUILDING CONTROL SUBMISSION READINESS",
]

PLANNING_REQUIRED_HEADINGS = [
    "PROJECT CLASSIFICATION",
    "SITE AND PROPOSAL OVERVIEW",
    "TOP SUMMARY",
    "LOCAL AUTHORITY CONTEXT",
    "PD / PRIOR APPROVAL / PLANNING ROUTE",
    "COMPLIANCE SNAPSHOT",
    "PLANNING ASSESSMENT",
    "DRAWING-PACK INCONSISTENCIES",
    "KEY RISKS",
    "MISSING INFORMATION",
    "RECOMMENDED ACTIONS",
    "PROFESSIONAL CONCLUSION",
    "SUBMISSION READINESS",
]

BUILDING_SECTION_ORDER = [
    ("PROJECT CLASSIFICATION", "Project Classification"),
    ("PROJECT DETAILS", "Project Details"),
    ("TOP SUMMARY", "Top Summary"),
    ("DRAWING-PACK INCONSISTENCIES", "Drawing-Pack Inconsistencies"),
    ("EXECUTIVE SUMMARY", "Executive Summary"),
    ("DRAWING PACK SUMMARY", "Drawing Pack Summary"),
    ("COMPLIANCE STATUS BY APPROVED DOCUMENT", "Compliance Status by Approved Document"),
    ("KEY RISKS", "Key Risks"),
    ("MISSING INFORMATION", "Missing Information"),
    ("RECOMMENDED ACTIONS", "Recommended Actions"),
    ("BUILDING CONTROL SUBMISSION READINESS", "Building Control Submission Readiness"),
]

PLANNING_SECTION_ORDER = [
    ("PROJECT CLASSIFICATION", "Project Classification"),
    ("SITE AND PROPOSAL OVERVIEW", "Site and Proposal Overview"),
    ("TOP SUMMARY", "Top Summary"),
    ("LOCAL AUTHORITY CONTEXT", "Local Authority Context"),
    ("PD / PRIOR APPROVAL / PLANNING ROUTE", "PD / Prior Approval / Planning Route"),
    ("COMPLIANCE SNAPSHOT", "Compliance Snapshot"),
    ("PLANNING ASSESSMENT", "Planning Assessment"),
    ("DRAWING-PACK INCONSISTENCIES", "Drawing-Pack Inconsistencies"),
    ("KEY RISKS", "Key Risks"),
    ("MISSING INFORMATION", "Missing Information"),
    ("RECOMMENDED ACTIONS", "Recommended Actions"),
    ("PROFESSIONAL CONCLUSION", "Professional Conclusion"),
    ("SUBMISSION READINESS", "Submission Readiness"),
]

BUILDING_SPECIAL_KEY_VALUE_SECTIONS = {
    "PROJECT CLASSIFICATION",
    "TOP SUMMARY",
    "BUILDING CONTROL SUBMISSION READINESS",
}

PLANNING_SPECIAL_KEY_VALUE_SECTIONS = {
    "PROJECT CLASSIFICATION",
    "TOP SUMMARY",
    "LOCAL AUTHORITY CONTEXT",
    "PD / PRIOR APPROVAL / PLANNING ROUTE",
    "COMPLIANCE SNAPSHOT",
    "SUBMISSION READINESS",
}

BUILDING_DISCLAIMER_TEXT = (
    "Beta: This report is an AI-assisted preliminary review. "
    "It does not replace professional Building Control approval, structural engineering design, "
    "or statutory review by the Local Authority or Approved Inspector."
)

PLANNING_DISCLAIMER_TEXT = (
    "This report provides an initial planning review based on the submitted drawings and information. "
    "It does not replace a formal planning appraisal or the decision of the Local Planning Authority."
)

MODULE_CONFIG = {
    "Building Regulations Review": {
        "required_headings": BUILDING_REQUIRED_HEADINGS,
        "section_order": BUILDING_SECTION_ORDER,
        "special_key_value_sections": BUILDING_SPECIAL_KEY_VALUE_SECTIONS,
        "disclaimer": BUILDING_DISCLAIMER_TEXT,
        "title": "Building Regulations Review Report",
        "readiness_key": "BUILDING CONTROL SUBMISSION READINESS",
    },
    "Planning Review": {
        "required_headings": PLANNING_REQUIRED_HEADINGS,
        "section_order": PLANNING_SECTION_ORDER,
        "special_key_value_sections": PLANNING_SPECIAL_KEY_VALUE_SECTIONS,
        "disclaimer": PLANNING_DISCLAIMER_TEXT,
        "title": "Planning Appraisal Report",
        "readiness_key": "SUBMISSION READINESS",
    },
}

PROJECT_TYPE_OPTIONS = [
    "Ground Floor Rear Extension",
    "Ground Floor Side Extension",
    "Ground Floor Infill Extension",
    "Porch",
    "First Floor Rear Extension",
    "First Floor Side Extension",
    "Loft Extension",
    "Flat Conversion",
    "House Conversion",
]

PROPERTY_TYPE_OPTIONS = [
    "Not stated",
    "Detached House",
    "Semi-Detached House",
    "Terraced House",
    "End of Terrace House",
    "Bungalow",
    "Chalet Bungalow",
    "Flat",
    "Maisonette",
    "Other",
]


CLASS_A_PROJECT_TYPES = {
    "Ground Floor Rear Extension",
    "Ground Floor Side Extension",
    "Ground Floor Infill Extension",
    "First Floor Rear Extension",
    "First Floor Side Extension",
}

SCOPE_ITEM_OPTIONS = [
    "New kitchen",
    "New bathroom / WC",
    "New bedroom",
    "Internal room layout changes",
    "Staircase changes",
    "Structural openings",
    "New steel beams",
    "Drainage changes",
    "New external doors / windows",
    "Roof / loft / dormer works",
    "Insulation / thermal upgrades",
    "Heating / boiler / ventilation",
    "Fire safety upgrades",
    "Sound insulation",
    "New foundations",
    "Basement excavation",
    "Party wall interface",
]



def get_accuracy_question_group(project_types: List[str]) -> str:
    selected = set(project_types or [])
    if "Loft Extension" in selected:
        return "class_b"
    if selected & CLASS_A_PROJECT_TYPES:
        return "class_a"
    if "Porch" in selected:
        return "class_d"
    return "generic"


def render_improve_accuracy_section(project_types: List[str]):
    answers: Dict[str, str] = {}
    with st.expander("Improve Accuracy (Optional)", expanded=False):
        st.caption("Simple extra questions based on the householder technical guidance. Only the relevant questions for the selected project type are shown.")
        client_name = st.text_input("Client")
        review_date = st.date_input("Report Date")
        group = get_accuracy_question_group(project_types)

        if group == "class_b":
            answers["pd_question_family"] = "class_b"
            st.caption("Class B – additions etc. to the roof")
            answers["site_constraints"] = st.multiselect(
                "Any site constraints that may affect PD rights?",
                ["None", "Conservation area / Article 2(3) land", "Article 4 direction", "Listed building"],
                default=["None"],
                key="acc_b_constraints",
            )
            answers["front_roof_plane_highway"] = st.selectbox(
                "Does any part project from the front roof slope facing the highway?",
                ["Not sure", "No", "Yes"],
                key="acc_b_front",
            )
            answers["roof_volume_band"] = st.selectbox(
                "Added roof volume",
                ["Not sure", "Within normal limit", "Over limit"],
                key="acc_b_volume",
                help="Use 'Within normal limit' where the total added roof volume appears to stay within the normal Class B allowance for the house type.",
            )
            answers["above_existing_roof_height"] = st.selectbox(
                "Does it rise above the highest part of the existing roof?",
                ["Not sure", "No", "Yes"],
                key="acc_b_highest",
            )
            answers["eaves_setback_0_2m"] = st.selectbox(
                "Is the enlargement set back at least 200mm from the original eaves?",
                ["Not sure", "Yes", "No"],
                key="acc_b_eaves",
            )
            answers["side_windows_obscure_glazed"] = st.selectbox(
                "Any side windows? If yes, will they be obscure glazed and non-opening below 1.7m?",
                ["Not applicable", "Yes", "No", "Not sure"],
                key="acc_b_windows",
            )
            answers["materials_similar"] = st.selectbox(
                "Will external materials be similar in appearance to the existing house?",
                ["Not sure", "Yes", "No"],
                key="acc_b_materials",
            )

        elif group == "class_a":
            answers["pd_question_family"] = "class_a"
            st.caption("Class A – enlargement, improvement or alteration")
            answers["site_constraints"] = st.multiselect(
                "Any site constraints that may affect PD rights?",
                ["None", "Conservation area / Article 2(3) land", "Article 4 direction", "Listed building"],
                default=["None"],
                key="acc_a_constraints",
            )
            answers["forward_of_principal_elevation"] = st.selectbox(
                "Does any part project beyond the principal elevation or a side elevation facing a highway?",
                ["Not sure", "No", "Yes"],
                key="acc_a_principal",
            )
            answers["existing_rear_extensions"] = st.selectbox(
                "Have there been previous rear extensions added to the original house?",
                ["Not sure", "No", "Yes"],
                key="acc_a_existing",
            )
            answers["within_2m_of_boundary"] = st.selectbox(
                "Is any part within 2m of a boundary?",
                ["Not sure", "No", "Yes"],
                key="acc_a_boundary",
            )
            answers["eaves_height_within_2m"] = st.selectbox(
                "If within 2m of a boundary, are the eaves 3.0m or lower?",
                ["Not applicable", "Yes", "No", "Not sure"],
                key="acc_a_eaves",
            )
            answers["side_extension_width"] = st.selectbox(
                "For a side extension, is it more than half the width of the original house?",
                ["Not applicable", "No", "Yes", "Not sure"],
                key="acc_a_width",
            )
            answers["materials_similar"] = st.selectbox(
                "Will external materials be similar in appearance to the existing house?",
                ["Not sure", "Yes", "No"],
                key="acc_a_materials",
            )

        elif group == "class_d":
            answers["pd_question_family"] = "class_d"
            st.caption("Class D – porch")
            answers["site_constraints"] = st.multiselect(
                "Any site constraints that may affect PD rights?",
                ["None", "Conservation area / Article 2(3) land", "Article 4 direction", "Listed building"],
                default=["None"],
                key="acc_d_constraints",
            )
            answers["porch_ground_area_band"] = st.selectbox(
                "Is the porch ground area 3m² or less?",
                ["Not sure", "Yes", "No"],
                key="acc_d_area",
            )
            answers["porch_height_band"] = st.selectbox(
                "Is the porch 3.0m high or lower?",
                ["Not sure", "Yes", "No"],
                key="acc_d_height",
            )
            answers["porch_within_2m_highway"] = st.selectbox(
                "Is any part within 2m of a boundary with a highway?",
                ["Not sure", "No", "Yes"],
                key="acc_d_highway",
            )
        else:
            st.caption("Add only anything that clearly changes the planning route.")

        answers["accuracy_notes"] = st.text_area(
            "Anything else that could affect PD, prior approval, or full planning route?",
            height=90,
            placeholder="Example: Article 4 area, listed building, previous extensions already built, or front dormer proposed.",
            key=f"acc_notes_{group}",
        )

    return client_name, review_date, answers


def get_planning_route_snapshot(
    project_types: List[str],
    property_type: str,
    proposal_summary: str,
    rear_extension_depth_m=None,
    rear_extension_height_m=None,
    accuracy_answers: Dict[str, str] | None = None,
):
    selected = set(project_types or [])
    accuracy_answers = accuracy_answers or {}
    summary = (proposal_summary or '').lower()
    property_type_l = (property_type or '').lower()

    route = 'Full Planning likely'
    risk = 'Medium'
    reason = 'Add core project details to improve route accuracy.'

    if 'Loft Extension' in selected:
        front = accuracy_answers.get('front_roof_plane_highway', 'Not sure')
        highest = accuracy_answers.get('above_existing_roof_height', 'Not sure')
        volume = accuracy_answers.get('roof_volume_band', 'Not sure')
        if front == 'No' and highest == 'No' and volume == 'Within normal limit':
            route = 'PD / LDC possible'
            risk = 'Medium'
            reason = 'Class B may apply if the roof enlargement stays off the front roof slope, stays below the highest roof part, and remains within the normal volume allowance.'
        elif front == 'Yes' or highest == 'Yes' or volume == 'Over limit':
            route = 'Full Planning likely'
            risk = 'High'
            reason = 'Class B is less likely where the loft addition affects the front roof slope, exceeds the highest roof part, or appears above the normal volume allowance.'
        else:
            route = 'PD / LDC possible'
            risk = 'Medium'
            reason = 'Loft works may fall under Class B, but the key front roof, highest roof, and volume checks still need confirmation.'

    elif selected & CLASS_A_PROJECT_TYPES:
        depth = rear_extension_depth_m or 0.0
        height = rear_extension_height_m or 0.0
        detached = 'detached' in property_type_l
        attached = any(x in property_type_l for x in ['semi', 'terraced', 'end of terrace'])
        principal = accuracy_answers.get('forward_of_principal_elevation', 'Not sure')
        within_boundary = accuracy_answers.get('within_2m_of_boundary', 'Not sure')
        eaves_test = accuracy_answers.get('eaves_height_within_2m', 'Not sure')

        if principal == 'Yes' or 'first floor' in summary or {'First Floor Rear Extension', 'First Floor Side Extension'} & selected:
            route = 'Full Planning likely'
            risk = 'High'
            reason = 'Upper-floor enlargements and works projecting beyond the principal elevation are usually outside straightforward Class A PD.'
        else:
            pd_limit = 4.0 if detached else 3.0
            pa_limit = 8.0 if detached else 6.0
            if 'Ground Floor Rear Extension' in selected and depth and depth <= pd_limit and height <= 4.0 and not (within_boundary == 'Yes' and eaves_test == 'No'):
                route = 'PD / LDC possible'
                risk = 'Low'
                reason = 'The extension appears within standard Class A depth and height limits, subject to full PD checks.'
            elif 'Ground Floor Rear Extension' in selected and depth and depth <= pa_limit and height <= 4.0:
                route = 'Prior Approval likely'
                risk = 'Medium'
                reason = 'The rear extension appears to fall within the larger home extension range, so the neighbour consultation / prior approval route may apply.'
            else:
                route = 'Full Planning likely'
                risk = 'High'
                reason = 'The extension appears to exceed the usual Class A / larger home extension limits or needs fuller planning assessment.'

    elif 'Porch' in selected:
        area = accuracy_answers.get('porch_ground_area_band', 'Not sure')
        h = accuracy_answers.get('porch_height_band', 'Not sure')
        highway = accuracy_answers.get('porch_within_2m_highway', 'Not sure')
        if area == 'Yes' and h == 'Yes' and highway == 'No':
            route = 'PD possible'
            risk = 'Low'
            reason = 'The porch appears capable of falling within Class D size, height, and highway-distance limits.'
        elif area == 'No' or h == 'No' or highway == 'Yes':
            route = 'Full Planning likely'
            risk = 'High'
            reason = 'The porch appears to fail one or more Class D limits on area, height, or distance to the highway boundary.'
        else:
            route = 'PD possible'
            risk = 'Medium'
            reason = 'A porch may fall under Class D, but the area, height, and highway checks still need confirmation.'

    elif 'Flat Conversion' in selected or 'House Conversion' in selected or 'flat' in property_type_l or 'maisonette' in property_type_l:
        route = 'Full Planning likely'
        risk = 'High'
        reason = 'Conversions and works to flats / maisonettes usually need fuller planning review rather than householder PD.'

    return route, risk, reason


def build_accuracy_context(answers: Dict[str, str]) -> str:
    if not answers:
        return ""
    label_map = {
        'site_constraints': 'Site constraints',
        'front_roof_plane_highway': 'Front roof slope facing highway',
        'roof_volume_band': 'Added roof volume',
        'above_existing_roof_height': 'Above highest roof',
        'eaves_setback_0_2m': '200mm eaves setback',
        'side_windows_obscure_glazed': 'Side windows obscure glazed',
        'materials_similar': 'Materials similar to existing house',
        'forward_of_principal_elevation': 'Projects beyond principal elevation / highway side',
        'existing_rear_extensions': 'Previous rear extensions to original house',
        'within_2m_of_boundary': 'Within 2m of boundary',
        'eaves_height_within_2m': 'Boundary eaves test',
        'side_extension_width': 'Side extension more than half width',
        'porch_ground_area_band': 'Porch ground area 3m² or less',
        'porch_height_band': 'Porch height 3m or lower',
        'porch_within_2m_highway': 'Porch within 2m of highway boundary',
        'accuracy_notes': 'Extra route notes',
    }
    lines = []
    for key, value in answers.items():
        if key == 'pd_question_family':
            continue
        if isinstance(value, list):
            value_str = ", ".join(v for v in value if v and v != 'None').strip()
        else:
            value_str = str(value).strip()
        if not value_str or value_str in {'Not sure', 'Not stated', '', 'None'}:
            continue
        lines.append(f"{label_map.get(key, key)}: {value_str}")
    return " | ".join(lines)


def build_pd_context(project_types: List[str], property_type: str, rear_extension_depth_m, rear_extension_height_m, accuracy_answers: Dict[str, str]) -> Dict[str, str]:
    pd_context: Dict[str, str] = dict(accuracy_answers or {})
    pd_context['is_single_dwellinghouse'] = 'no' if (property_type or '').strip().lower() in {'flat', 'maisonette'} else 'yes'
    constraints = pd_context.get('site_constraints', [])
    if isinstance(constraints, list):
        pd_context['site_constraints'] = ', '.join([c for c in constraints if c and c != 'None']) or 'None'
    if rear_extension_depth_m is not None:
        pd_context['rear_extension_depth_m'] = f"{float(rear_extension_depth_m):.1f}"
    if rear_extension_height_m is not None:
        pd_context['rear_extension_overall_height_m'] = f"{float(rear_extension_height_m):.1f}"
    if not pd_context.get('pd_question_family'):
        pd_context['pd_question_family'] = get_accuracy_question_group(project_types)
    return pd_context





# -----------------------------------------------------------------------------
# AI CONFIDENCE SYSTEM
# Rule-engine-first confidence labels. These are not legal certainty scores.
# They are user-facing status labels based on deterministic checks, drawing
# completeness, report sections and missing information.
# -----------------------------------------------------------------------------
PLANNING_CONFIDENCE_LABELS = [
    "LIKELY COMPLIANT",
    "LIKELY COMPLIANT SUBJECT TO MINOR CHECKS",
    "LIKELY PRIOR APPROVAL",
    "REQUIRES FURTHER REVIEW",
    "LIKELY PLANNING PERMISSION REQUIRED",
]

BUILDING_CONFIDENCE_LABELS = [
    "LIKELY COMPLIANT",
    "PARTIAL INFORMATION",
    "NON-COMPLIANT ITEMS FOUND",
    "STRUCTURAL REVIEW REQUIRED",
    "FIRE STRATEGY REVIEW REQUIRED",
    "BUILDING CONTROL REVIEW ADVISED",
]


def _normalise_text(value) -> str:
    return str(value or "").strip()




def _is_minor_class_b_condition_issue_text(text_value: str) -> bool:
    """Detect Class B/C loft cases where the only issue is a side-window
    obscurity/non-opening annotation. This should not be presented as full
    planning required where the main PD/LDC route remains available.
    """
    text = _normalise_text(text_value).upper()
    if not any(t in text for t in ["CLASS B", "DORMER", "LOFT", "ROOF ENLARGEMENT", "ROOFLIGHT"]):
        return False
    if not any(t in text for t in ["SIDE WINDOWS", "SIDE-FACING", "SIDE ROOF WINDOWS", "OBSCURE", "1.7M"]):
        return False
    hard_fail_terms = [
        "FRONT-FACING ROOF ENLARGEMENT",
        "PRINCIPAL ELEVATION AND FRONTS A HIGHWAY",
        "ABOVE THE HIGHEST PART",
        "EXCEEDS HIGHEST ROOF",
        "OVER LIMIT",
        "ROOF VOLUME EXCEEDS",
        "BALCONY",
        "VERANDAH",
        "RAISED PLATFORM",
        "ARTICLE 4",
        "LISTED BUILDING",
        "FLAT OR MAISONETTE",
        "NOT A SINGLE DWELLINGHOUSE",
    ]
    return not any(term in text for term in hard_fail_terms)

def _extract_route_from_rule_summary(rule_summary: str) -> str:
    """Pull a stable route/status label from the deterministic rule summary.

    Planning reports should use professional status wording rather than hard PASS/FAIL labels.
    Missing minor confirmations should not create a fail result for otherwise typical PD schemes.
    """
    text = _normalise_text(rule_summary).upper()
    if _is_minor_class_b_condition_issue_text(text):
        return "LIKELY COMPLIANT SUBJECT TO MINOR CHECKS"
    if "PRIOR APPROVAL" in text:
        return "LIKELY PRIOR APPROVAL"
    if "FULL PLANNING" in text or "PLANNING PERMISSION" in text:
        return "LIKELY PLANNING PERMISSION REQUIRED"
    if "PD / LDC" in text or "PERMITTED DEVELOPMENT" in text or "PD POSSIBLE" in text or "LIKELY PD" in text or "PASS" in text:
        return "LIKELY COMPLIANT"
    if "FAIL" in text:
        return "LIKELY PLANNING PERMISSION REQUIRED"
    if "NEEDS CONFIRMATION" in text or "MANUAL" in text:
        return "REQUIRES FURTHER REVIEW"
    return "REQUIRES FURTHER REVIEW"

def _count_report_signals(sections: Dict[str, str], needles: List[str]) -> int:
    combined = "\n".join(sections.values()).upper()
    return sum(1 for n in needles if n.upper() in combined)


def calculate_planning_confidence(sections: Dict[str, str], rule_summary: str = "") -> Dict[str, object]:
    """Return user-facing planning confidence label and explanation.

    This avoids fake percentage certainty and avoids harsh FAIL/PASS wording.
    The label is based on the deterministic rule-engine result first, then adjusted
    for actual report findings. Typical PD/LDC schemes with minor missing checks
    should show as likely compliant subject to minor checks, not failed.
    """
    route_label = _extract_route_from_rule_summary(rule_summary)
    missing_text = _normalise_text(sections.get("MISSING INFORMATION", "")).upper()
    risk_text = _normalise_text(sections.get("KEY RISKS", "")).upper()
    readiness_text = _normalise_text(sections.get("SUBMISSION READINESS", "")).upper()
    route_text = _normalise_text(sections.get("PD / PRIOR APPROVAL / PLANNING ROUTE", "")).upper()
    top_text = _normalise_text(sections.get("TOP SUMMARY", "")).upper()
    combined = "\n".join([missing_text, risk_text, readiness_text, route_text, top_text, _normalise_text(rule_summary).upper()])

    blockers = _count_report_signals(sections, [
        "NOT CLEARLY SHOWN",
        "NOT CLEARLY DIMENSIONED",
        "INSUFFICIENT",
        "REQUIRES CONFIRMATION",
        "MISSING",
    ])
    minor_class_b_condition_only = _is_minor_class_b_condition_issue_text(combined)
    clear_policy_issue = any(x in combined for x in [
        "FULL PLANNING REQUIRED",
        "LIKELY PLANNING PERMISSION REQUIRED",
        "EXCEEDS",
        "OUTSIDE CLASS",
        "FRONT-FACING ROOF ENLARGEMENT",
        "ABOVE THE HIGHEST PART",
        "ARTICLE 4",
        "LISTED BUILDING",
    ])
    if minor_class_b_condition_only:
        clear_policy_issue = False
    high_risk = "HIGH" in risk_text or "NOT READY" in readiness_text

    if minor_class_b_condition_only:
        label = "LIKELY COMPLIANT SUBJECT TO MINOR CHECKS"
    elif route_label == "LIKELY PRIOR APPROVAL" and not clear_policy_issue:
        label = "LIKELY PRIOR APPROVAL"
    elif route_label == "LIKELY PLANNING PERMISSION REQUIRED" or clear_policy_issue:
        label = "LIKELY PLANNING PERMISSION REQUIRED"
    elif route_label == "LIKELY COMPLIANT":
        label = "LIKELY COMPLIANT" if blockers <= 1 and not high_risk else "LIKELY COMPLIANT SUBJECT TO MINOR CHECKS"
    elif route_label == "LIKELY COMPLIANT SUBJECT TO MINOR CHECKS":
        label = "LIKELY COMPLIANT SUBJECT TO MINOR CHECKS"
    elif blockers <= 3 and ("PD / LDC" in combined or "CLASS B" in combined or "PERMITTED DEVELOPMENT" in combined):
        label = "LIKELY COMPLIANT SUBJECT TO MINOR CHECKS"
    else:
        label = "REQUIRES FURTHER REVIEW"

    triggers = []
    if minor_class_b_condition_only:
        triggers.append("Loft/dormer proposal appears capable of a PD/LDC route, subject to side-window annotation checks.")
    elif "CLASS B" in combined or "DORMER" in combined or "ROOF" in combined:
        triggers.append("Roof enlargement checks appear to be the main planning route issue.")
    elif "PRIOR APPROVAL" in combined:
        triggers.append("Larger home extension / prior approval route appears relevant.")
    elif "FULL PLANNING" in combined:
        triggers.append("The proposal may require a householder planning application.")
    else:
        triggers.append("Planning route has been assessed using the uploaded drawings and rule checks.")
    if blockers:
        triggers.append("Some standard confirmation items may still need checking before submission.")
    if "PD / LDC" in combined or "PERMITTED DEVELOPMENT" in combined:
        triggers.append("A Lawful Development Certificate route may be appropriate where PD criteria are met.")

    return {
        "module": "Planning Review",
        "label": label,
        "basis": "Planning rules + drawing review",
        "triggers": triggers[:3],
        "note": "Indicative review only. The local authority makes the final decision.",
    }

def calculate_building_confidence(sections: Dict[str, str]) -> Dict[str, object]:
    combined = "\n".join(sections.values()).upper()
    compliance = _normalise_text(sections.get("COMPLIANCE STATUS BY APPROVED DOCUMENT", "")).upper()
    missing = _normalise_text(sections.get("MISSING INFORMATION", "")).upper()
    risks = _normalise_text(sections.get("KEY RISKS", "")).upper()

    fail_count = compliance.count("FAIL") + combined.count("NON-COMPLIANT")
    partial_count = compliance.count("PARTLY") + compliance.count("REVIEW REQUIRED") + missing.count("NOT CLEARLY")

    if "STRUCTURAL" in risks and any(x in risks for x in ["REQUIRED", "CALC", "ENGINEER"]):
        label = "STRUCTURAL REVIEW REQUIRED"
    elif "FIRE" in risks and any(x in risks for x in ["UNCLEAR", "REQUIRED", "ESCAPE", "STRATEGY"]):
        label = "FIRE STRATEGY REVIEW REQUIRED"
    elif fail_count > 0:
        label = "NON-COMPLIANT ITEMS FOUND"
    elif partial_count > 0 or "MISSING" in missing:
        label = "PARTIAL INFORMATION"
    elif "READY" in _normalise_text(sections.get("BUILDING CONTROL SUBMISSION READINESS", "")).upper():
        label = "LIKELY COMPLIANT"
    else:
        label = "BUILDING CONTROL REVIEW ADVISED"

    triggers = []
    if fail_count:
        triggers.append("One or more Approved Document checks are marked as fail/non-compliant.")
    if partial_count:
        triggers.append("Some compliance items require further details or confirmation.")
    if "STRUCTURAL" in combined:
        triggers.append("Structural information or engineer confirmation is relevant.")
    if "FIRE" in combined:
        triggers.append("Fire strategy information is relevant to the review.")
    if not triggers:
        triggers.append("No major issue was identified from the visible report sections.")

    return {
        "module": "Building Regulations Review",
        "label": label,
        "basis": "Drawing/report checks",
        "triggers": triggers[:4],
        "note": "This is an AI-assisted review status. Formal Building Control approval is still required.",
    }


def calculate_ai_confidence(module_name: str, sections: Dict[str, str], rule_summary: str = "") -> Dict[str, object]:
    if module_name == "Planning Review":
        return calculate_planning_confidence(sections, rule_summary)
    return calculate_building_confidence(sections)


def confidence_badge_style(label: str) -> str:
    label_u = _normalise_text(label).upper()
    if label_u in {"LIKELY COMPLIANT", "LIKELY COMPLIANT SUBJECT TO MINOR CHECKS", "LIKELY PRIOR APPROVAL"}:
        return "background:#DDF3E4;color:#14532D;border-color:#B7E4C7;"
    if label_u in {"LIKELY PLANNING PERMISSION REQUIRED", "REQUIRES FURTHER REVIEW", "PARTIAL INFORMATION", "STRUCTURAL REVIEW REQUIRED", "FIRE STRATEGY REVIEW REQUIRED", "BUILDING CONTROL REVIEW ADVISED"}:
        return "background:#FFF3CD;color:#6B4E00;border-color:#F1D48A;"
    return "background:#F8D7DA;color:#842029;border-color:#F1AEB5;"


def render_ai_confidence_card(confidence):
    if not confidence:
        return
    label = _normalise_text(confidence.get("label", "MANUAL REVIEW ADVISED"))
    style = confidence_badge_style(label)
    triggers = confidence.get("triggers", []) or []
    st.markdown(
        f"""
        <div class="sy-subtle-card">
            <div class="sy-section-label">Planning Confidence</div>
            <div style="display:flex;align-items:center;justify-content:space-between;gap:1rem;flex-wrap:wrap;">
                <h3 style="margin:0;">Review Status</h3>
                <span style="display:inline-block;padding:0.42rem 0.72rem;border-radius:999px;border:1px solid; font-weight:800; letter-spacing:0.02em; {style}">{label}</span>
            </div>
            <div class="sy-muted" style="margin-top:0.45rem;">Basis: {confidence.get('basis', 'Rule and drawing checks')}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if triggers:
        st.markdown("**What caused this result?**")
        for t in triggers:
            st.markdown(f"- {t}")
    if confidence.get("note"):
        st.caption(str(confidence.get("note")))

PLAN_LABELS = {
    "starter": "Solo",
    "pro": "Studio",
}
WEBSITE_PRICING_URL = "https://www.sydesignstudio.co.uk/pricing-plans"
WEBSITE_HOME_URL = "https://www.sydesignstudio.co.uk"
WEBSITE_LAUNCH_URL = "https://www.sydesignstudio.co.uk/archlens-launch"
ARCHLENS_SHARED_SECRET = "ArchLens-SYDS-2026-very-long-random-private-secret-839201"

def get_verified_plan_and_user() -> Tuple[str, str, bool]:
    try:
        query_params = st.query_params
        token_value = query_params.get("token", "")
    except Exception:
        token_value = ""

    if isinstance(token_value, list):
        token_value = token_value[0] if token_value else ""

    token_value = str(token_value).strip()
    if not token_value:
        return "starter", "", False

    try:
        payload = jwt.decode(
            token_value,
            ARCHLENS_SHARED_SECRET,
            algorithms=["HS256"],
        )
        plan = str(payload.get("plan", "starter")).strip().lower()
        if plan not in {"starter", "pro"}:
            plan = "starter"

        email = str(payload.get("email", "")).strip()
        display_name = email or str(payload.get("sub", "")).strip()
        return plan, display_name, True
    except Exception:
        return "starter", "", False


def get_allowed_review_modules(plan: str) -> List[str]:
    # Token system: both modules can be selected by authenticated users.
    # Downloads are controlled by credits rather than the old monthly plan gate.
    return ["Planning Review", "Building Regulations Review"]


def get_plan_upgrade_message(feature_name: str) -> str:
    return f"{feature_name} is available on Studio. Upgrade to unlock this feature."

def add_saved_project(project_record: Dict):
    saved = st.session_state.get("saved_projects", [])
    filtered = [item for item in saved if item.get("report_id") != project_record.get("report_id")]
    filtered.insert(0, project_record)
    st.session_state["saved_projects"] = filtered[:25]


def get_credit_balance() -> int:
    return int(st.session_state.get("credit_balance", 0) or 0)


def add_credit_transaction(amount: int, reason: str, report_id: str = "", balance_after=None):
    transactions = st.session_state.get("credit_transactions", []) or []
    transactions.insert(0, {
        "date": time.strftime("%Y-%m-%d %H:%M"),
        "amount": int(amount),
        "reason": reason,
        "report_id": report_id,
        "balance_after": get_credit_balance() if balance_after is None else balance_after,
    })
    st.session_state["credit_transactions"] = transactions[:100]


def grant_credits(amount: int, reason: str = "Credits added"):
    new_balance = get_credit_balance() + int(amount)
    st.session_state["credit_balance"] = new_balance
    add_credit_transaction(int(amount), reason, balance_after=new_balance)


def spend_credits(amount: int, reason: str, report_id: str = "", export_type: str = ""):
    amount = int(amount)
    balance = get_credit_balance()

    if amount <= 0:
        return True, "No credits required."

    if balance < amount:
        return False, f"Not enough credits. You need {amount} credits but only have {balance}."

    user_email = normalise_user_email(st.session_state.get("auth_user_name", ""))

    # Live persistence: deduct from ArchLens API first, then update Streamlit session.
    api_result = api_deduct_credits(
        user_email,
        amount,
        report_id=report_id,
        export_type=export_type,
    )

    if not api_result.get("success"):
        return False, api_result.get("message", "Credit deduction failed. Please try again.")

    new_balance = int(api_result.get("credits", max(0, balance - amount)) or 0)
    st.session_state["credit_balance"] = new_balance
    add_credit_transaction(-amount, reason, report_id=report_id, balance_after=new_balance)
    return True, f"Unlocked successfully. {amount} credits used. New balance: {new_balance}."


def get_export_credit_cost(module_name: str, export_type: str) -> int:
    return int(EXPORT_CREDIT_COSTS.get(module_name, {}).get(export_type, 0))


def is_report_unlocked(report_id: str, export_type: str) -> bool:
    unlocked = st.session_state.get("unlocked_reports", {}) or {}
    return bool(unlocked.get(report_id, {}).get(export_type, False))


def mark_report_unlocked(report_id: str, export_type: str, cost: int = 0):
    unlocked = st.session_state.get("unlocked_reports", {}) or {}
    unlocked.setdefault(report_id, {})[export_type] = True
    st.session_state["unlocked_reports"] = unlocked
    saved = st.session_state.get("saved_projects", []) or []
    for item in saved:
        if item.get("report_id") == report_id:
            item[f"{export_type}_unlocked"] = True
            item["credits_used"] = int(item.get("credits_used", 0) or 0) + int(cost or 0)
    st.session_state["saved_projects"] = saved


def unlock_report_export(report_id: str, module_name: str, export_type: str):
    if is_report_unlocked(report_id, export_type):
        return True, "Already unlocked. You can download again without using more credits."
    cost = get_export_credit_cost(module_name, export_type)
    ok, message = spend_credits(cost, f"Unlock {export_type.upper()} export", report_id=report_id, export_type=export_type)
    if ok:
        mark_report_unlocked(report_id, export_type, cost)
    return ok, message


def render_credit_balance_card(compact: bool = False):
    balance = get_credit_balance()
    if compact:
        st.caption(f"Credits: {balance}")
        return
    st.markdown(
        f'''<div class="sy-subtle-card">
            <div class="sy-section-label">Credits</div>
            <h3 style="margin:0;">{balance} credits available</h3>
            <div class="sy-muted" style="margin-top:0.35rem;">{FREE_PREVIEW_NOTE}</div>
        </div>''',
        unsafe_allow_html=True,
    )


def render_buy_credits_panel():
    st.markdown("### Buy Credits")
    st.caption("Stripe/Wix payment automation is the next phase. These buttons simulate credit top-ups for testing the app flow.")
    cols = st.columns(len(CREDIT_PACKS))
    for idx, (pack_name, pack) in enumerate(CREDIT_PACKS.items()):
        with cols[idx]:
            st.markdown(f"**{pack_name}**")
            st.caption(pack["price"])
            if st.button(f"Add {pack['credits']} credits", key=f"add_credit_pack_{idx}", use_container_width=True):
                grant_credits(pack["credits"], f"Test top-up: {pack_name}")
                st.success(f"Added {pack['credits']} credits.")
                st.rerun()


def inject_custom_css():
    theme = st.session_state.get("app_theme", "Dark")
    light_mode = str(theme).lower().startswith("light")
    if light_mode:
        bg = "#F7F6F2"
        surface = "#FFFFFF"
        surface_2 = "#F1EFE8"
        text = "#171717"
        muted = "#5F6368"
        border = "#D9D6CC"
        shadow = "0 12px 28px rgba(20,20,20,0.08)"
        sidebar_bg = "#FFFFFF"
        input_bg = "#FFFFFF"
    else:
        bg = "#0E1117"
        surface = "#121821"
        surface_2 = "#172033"
        text = "#F5F7FA"
        muted = "#B8C0CC"
        border = "#2A3140"
        shadow = "0 14px 32px rgba(0,0,0,0.24)"
        sidebar_bg = "#0B0F16"
        input_bg = "#111827"

    st.markdown(
        f"""
        <style>
        :root {{
            --sy-bg: {bg};
            --sy-surface: {surface};
            --sy-surface-2: {surface_2};
            --sy-border: {border};
            --sy-text: {text};
            --sy-muted: {muted};
            --sy-accent: #D4C29A;
            --sy-accent-hover: #C5B183;
            --sy-card-shadow: {shadow};
            --sy-input-bg: {input_bg};
            --sy-sidebar-bg: {sidebar_bg};
        }}

        .stApp {{
            background: var(--sy-bg) !important;
            color: var(--sy-text) !important;
            font-size: 15px;
        }}
        header[data-testid="stHeader"], [data-testid="stToolbar"], .stAppDeployButton {{ display:none !important; }}
        #MainMenu {{ visibility:hidden !important; }}
        footer {{ visibility:hidden !important; }}

        .block-container {{
            padding-top: 1rem !important;
            padding-bottom: 2rem !important;
            max-width: 1420px !important;
        }}

        [data-testid="stSidebar"] {{
            background: var(--sy-sidebar-bg) !important;
            border-right: 1px solid var(--sy-border);
        }}
        [data-testid="stSidebar"] * {{ color: var(--sy-text); }}
        [data-testid="stSidebar"] [role="radiogroup"] label {{
            padding: 0.42rem 0.55rem !important;
            border-radius: 12px !important;
            margin-bottom: 0.15rem !important;
        }}

        h1 {{ font-size: 2.35rem !important; letter-spacing: -0.035em !important; line-height: 1.08 !important; }}
        h2 {{ font-size: 1.45rem !important; letter-spacing: -0.02em !important; }}
        h3 {{ font-size: 1.12rem !important; }}
        p, li, label, .stMarkdown, .stCaption {{ color: var(--sy-text); }}
        small, .sy-muted {{ color: var(--sy-muted) !important; }}

        .sy-topbar {{
            display:flex; justify-content:space-between; align-items:center; gap: 1rem;
            padding:0.85rem 1rem; border:1px solid var(--sy-border); border-radius:18px;
            background: var(--sy-surface); margin-bottom:1rem; box-shadow: var(--sy-card-shadow);
        }}
        .sy-topbar-title {{ font-size:0.78rem; text-transform:uppercase; letter-spacing:0.14em; color:var(--sy-muted); font-weight:800; }}
        .sy-topbar-meta {{ font-size:0.86rem; color:var(--sy-text); }}

        .sy-hero {{
            padding:1.25rem 1.3rem; border:1px solid var(--sy-border); border-radius:24px;
            background: var(--sy-surface); margin-bottom:1rem; box-shadow: var(--sy-card-shadow);
        }}
        .sy-hero-copy h1 {{ margin:0 0 0.45rem 0 !important; }}
        .sy-hero-copy .sy-muted {{ line-height:1.55; font-size:0.92rem; }}

        .sy-card, .sy-mini-card, .sy-upload-item, .sy-sidepanel, .sy-workspace, .sy-subtle-card, .sy-option-card, .sy-report-card {{
            border:1px solid var(--sy-border); background: var(--sy-surface); box-shadow: var(--sy-card-shadow); color: var(--sy-text);
        }}
        .sy-card {{ border-radius:20px; padding:1.05rem; margin-bottom:0.95rem; }}
        .sy-mini-card {{ border-radius:18px; padding:0.9rem; min-height:118px; }}
        .sy-sidepanel {{ border-radius:20px; padding:0.85rem; position:sticky; top:1rem; font-size:0.86rem; }}
        .sy-sidepanel .sy-data-row {{ font-size:0.82rem; padding:0.42rem 0; }}
        .sy-workspace {{ border-radius:22px; padding:1rem; }}
        .sy-subtle-card {{ border-radius:18px; padding:0.95rem 1rem; margin-bottom:0.85rem; }}
        .sy-section-label, .sy-panel-title, .sy-kpi {{
            font-size:0.72rem; text-transform:uppercase; letter-spacing:0.12em; color:var(--sy-muted); margin-bottom:0.35rem; font-weight:800;
        }}
        .sy-data-row {{ display:flex; justify-content:space-between; gap:0.8rem; padding:0.52rem 0; border-bottom:1px solid var(--sy-border); color:var(--sy-text); }}
        .sy-data-row:last-child {{ border-bottom:0; }}
        .sy-data-row span:first-child {{ color:var(--sy-muted); }}
        .sy-data-row strong {{ text-align:right; }}

        .sy-option-grid {{ display:grid; grid-template-columns: repeat(auto-fit, minmax(220px, 1fr)); gap:0.7rem; margin:0.75rem 0 1rem 0; }}
        .sy-option-card {{ border-radius:16px; padding:0.65rem 0.8rem; min-height:48px; }}
        .sy-option-card label {{ font-weight:650 !important; }}
        .sy-report-card {{ border-radius:18px; padding:1rem; margin-bottom:0.8rem; }}

        .sy-preview-shell {{ border:1px solid var(--sy-border); border-radius:18px; overflow:hidden; background: var(--sy-surface); }}
        .sy-preview-topbar {{ display:flex; justify-content:space-between; align-items:center; padding:0.75rem 0.9rem; border-bottom:1px solid var(--sy-border); background: var(--sy-surface-2); }}
        .sy-preview-title {{ font-weight:700; font-size:0.94rem; color:var(--sy-text); }}
        .sy-preview-meta {{ font-size:0.8rem; color:var(--sy-muted); }}
        .sy-preview-badge {{ padding:0.25rem 0.5rem; border-radius:999px; font-size:0.72rem; font-weight:700; background:rgba(212,194,154,0.18); color:var(--sy-text); border:1px solid var(--sy-border); }}
        .sy-preview-frame {{ border:0; background:white; }}
        .sy-empty-preview {{ min-height:240px; display:flex; align-items:center; justify-content:center; text-align:center; border:1px dashed var(--sy-border); border-radius:18px; background:var(--sy-surface-2); padding:1rem; color:var(--sy-muted); }}
        .sy-upload-item {{ border-radius:14px; padding:0.72rem 0.85rem; margin-bottom:0.5rem; }}

        div[data-testid="stMetric"] {{ background: var(--sy-surface); border:1px solid var(--sy-border); padding:0.72rem 0.85rem; border-radius:16px; }}
        div[data-testid="stMetric"] * {{ color: var(--sy-text) !important; }}

        .stDownloadButton button, .stButton button, .stLinkButton a {{ border-radius:14px !important; }}
        .stButton button, .stDownloadButton button, .stLinkButton a {{
            background: var(--sy-accent) !important; color: #111111 !important; border: 1px solid var(--sy-accent) !important;
            box-shadow: 0 10px 24px rgba(212, 194, 154, 0.18) !important; font-weight: 650 !important;
        }}
        .stButton button:hover, .stDownloadButton button:hover, .stLinkButton a:hover {{ background: var(--sy-accent-hover) !important; border-color: var(--sy-accent-hover) !important; color: #111111 !important; filter:none !important; }}

        .stSelectbox label, .stTextInput label, .stTextArea label, .stNumberInput label, .stDateInput label, .stMultiSelect label, .stCheckbox label, .stRadio label {{
            color:var(--sy-text) !important; font-weight:650 !important;
        }}
        [data-baseweb="select"] > div, [data-baseweb="tag"] {{ background:var(--sy-input-bg) !important; border:1px solid #5D6472 !important; color:var(--sy-text) !important; }}
        .stTextInput input, .stTextArea textarea, .stNumberInput input, .stDateInput input, div[data-baseweb="base-input"] > input, div[data-baseweb="base-input"] > textarea {{
            background:var(--sy-input-bg) !important; border:1px solid #5D6472 !important; color:var(--sy-text) !important; border-radius:12px !important;
        }}
        .stTextInput input:focus, .stTextArea textarea:focus, .stNumberInput input:focus, .stDateInput input:focus, div[data-baseweb="base-input"] > input:focus, div[data-baseweb="base-input"] > textarea:focus, [data-baseweb="select"] > div:focus-within {{
            border-color:#D4C29A !important; box-shadow:0 0 0 1px #D4C29A !important;
        }}
        .stTextArea textarea {{ min-height: 90px; }}
        .streamlit-expanderHeader {{ border:1px solid #5D6472 !important; border-radius:12px !important; }}
        .stProgress > div > div > div > div {{ background: linear-gradient(90deg, #D4C29A, #c5b183); }}
        </style>
        """,
        unsafe_allow_html=True,
    )

def smooth_progress(progress_bar, status_text, start, end, message, duration=0.8):
    steps = max(1, end - start)
    sleep_time = duration / steps
    for value in range(start, end + 1):
        progress_bar.progress(value)
        status_text.text(f"{message} {value}%")
        time.sleep(sleep_time)


def clean_input_value(value, fallback):
    if value is None:
        return fallback
    cleaned = str(value).strip()
    bad_values = {"", "a", "aa", "as", "s", "sa", "test", "xx", "ww", "w", "m", "na", "n/a"}
    if cleaned.lower() in bad_values:
        return fallback
    return cleaned




def extract_address_from_report(report_text: str, fallback: str = "Not provided") -> str:
    if not report_text:
        return fallback
    m = re.search(r"Project Address:\s*(.+)", report_text, re.IGNORECASE)
    if m:
        value = m.group(1).strip()
        if value and value.lower() != "not provided":
            return value
    m2 = re.search(r"\b\d+[A-Za-z]?\s+[^\n,]+(?:,\s*[^\n,]+){0,3},?\s*[A-Z]{1,2}\d[A-Z\d]?\s*\d[A-Z]{2}\b", report_text)
    if m2:
        return m2.group(0).strip()
    return fallback

def parse_key_value_lines(section_text: str) -> List[Tuple[str, str]]:
    rows = []
    for raw_line in section_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("- "):
            line = line[2:].strip()
        if ":" in line:
            label, value = line.split(":", 1)
            rows.append((label.strip(), value.strip()))
        else:
            rows.append(("", line))
    return rows


def parse_compliance_rows(content: str) -> List[Dict[str, str]]:
    rows = []
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("- "):
            line = line[2:].strip()
        if line.startswith("Part "):
            why_text = ""
            if ":" in line:
                left, status = line.split(":", 1)
                status = status.strip()
            else:
                left, status = line, ""

            if "–" in left:
                part_code, part_title = left.split("–", 1)
            elif "-" in left:
                part_code, part_title = left.split("-", 1)
            else:
                part_code, part_title = left, ""

            if i + 1 < len(lines):
                next_line = lines[i + 1]
                if next_line.startswith("- "):
                    next_line = next_line[2:].strip()
                if next_line.lower().startswith("why:"):
                    why_text = next_line[4:].strip()
                    i += 1

            rows.append(
                {
                    "part": part_code.replace("Part", "").strip(),
                    "title": part_title.strip(),
                    "status": status,
                    "why": why_text,
                }
            )
        i += 1
    return rows


def parse_report_sections(report_text: str, required_headings: List[str]) -> Dict[str, str]:
    headings = set(required_headings)
    sections: Dict[str, str] = {}
    current_heading = None
    current_lines: List[str] = []

    for line in report_text.splitlines():
        stripped = line.strip()
        if stripped.upper() in headings:
            if current_heading:
                sections[current_heading] = "\n".join(current_lines).strip()
            current_heading = stripped.upper()
            current_lines = []
        else:
            if current_heading:
                current_lines.append(line)

    if current_heading:
        sections[current_heading] = "\n".join(current_lines).strip()
    return sections


def validate_report_headings(report_text: str, required_headings: List[str]) -> Tuple[bool, List[str]]:
    report_upper = report_text.upper()
    missing = [heading for heading in required_headings if heading not in report_upper]
    return len(missing) == 0, missing


def get_pdf_page_count(pdf_path: str) -> int:
    doc = fitz.open(pdf_path)
    try:
        return len(doc)
    finally:
        doc.close()


def render_pdf_preview(uploaded_files):
    if not uploaded_files:
        st.markdown('<div class="sy-empty-preview">Upload a PDF to preview the drawing pack workspace.</div>', unsafe_allow_html=True)
        return

    selected_file = uploaded_files[-1]
    try:
        pdf_bytes = selected_file.getvalue()
        pdf_b64 = base64.b64encode(pdf_bytes).decode("utf-8")
        iframe = f'''
        <div class="sy-preview-shell">
            <div class="sy-preview-topbar">
                <div>
                    <div class="sy-preview-title">{selected_file.name}</div>
                    <div class="sy-preview-meta">{round(selected_file.size / (1024 * 1024), 2)} MB • Live drawing preview</div>
                </div>
                <div class="sy-preview-badge">PDF</div>
            </div>
            <iframe src="data:application/pdf;base64,{pdf_b64}" width="100%" height="760" type="application/pdf" class="sy-preview-frame"></iframe>
        </div>
        '''
        st.markdown(iframe, unsafe_allow_html=True)
    except Exception:
        st.markdown('<div class="sy-empty-preview">Preview not available for this file in the live browser view.</div>', unsafe_allow_html=True)

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, font_size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)


def set_table_col_widths(table, widths_inches):
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_word_header(section, title_text="ArchLens AI", subtitle_text="AI Review"):
    header = section.header
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = para.add_run(title_text + "\n")
    run1.bold = True
    run1.font.size = Pt(11)
    run2 = para.add_run(subtitle_text)
    run2.font.size = Pt(9)


def add_word_footer(section, practice_name="ArchLens AI", report_title="AI Review"):
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"{practice_name} | {report_title} | Page ")
    run.font.size = Pt(9)
    add_page_number(para)





# -----------------------------------------------------------------------------
# SY DESIGN STUDIO BRANDED REPORT EXPORT SETTINGS
# Place your logo in one of these paths in your Render/GitHub app:
#   assets/sy_design_studio_logo.png
#   assets/logo.png
#   sy_design_studio_logo.png
# The PDF will still export cleanly if no logo file is found.
# -----------------------------------------------------------------------------
SY_BRAND = {
    "practice_name": "SY Design Studio",
    "product_name": "ArchLens AI",
    "charcoal": "#333333",
    "gold": "#D4C29A",
    "light_gold": "#F4EFE3",
    "soft_grey": "#F6F6F4",
    "mid_grey": "#707070",
    "line_grey": "#D9D9D4",
}


def get_brand_logo_path():
    candidates = [
        "assets/sy_design_studio_logo.png",
        "assets/sy_design_studio_logo.jpg",
        "assets/logo.png",
        "assets/logo.jpg",
        "sy_design_studio_logo.png",
        "sy_design_studio_logo.jpg",
        "logo.png",
        "logo.jpg",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            return candidate
    return None

def build_pdf_report(file_name, address, client, date, practice_name, report_id, sections, module_name):
    """Create a polished SY Design Studio branded PDF for Planning and Building Regulations modules."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader

    config = MODULE_CONFIG[module_name]
    section_order = config["section_order"]
    special_key_value_sections = config["special_key_value_sections"]
    disclaimer_text = config["disclaimer"]
    report_title = config["title"]

    brand_practice = practice_name or SY_BRAND["practice_name"]
    logo_path = get_brand_logo_path()
    try:
        logo_bytes = st.session_state.get("brand_logo_bytes")
    except Exception:
        logo_bytes = None

    charcoal = colors.HexColor(SY_BRAND["charcoal"])
    gold = colors.HexColor(SY_BRAND["gold"])
    light_gold = colors.HexColor(SY_BRAND["light_gold"])
    soft_grey = colors.HexColor(SY_BRAND["soft_grey"])
    mid_grey = colors.HexColor(SY_BRAND["mid_grey"])
    line_grey = colors.HexColor(SY_BRAND["line_grey"])

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    left_margin = 48
    right_margin = 48
    top_margin = 58
    bottom_margin = 48
    usable_width = width - left_margin - right_margin
    y = height - top_margin
    page_no = 0

    def safe_text(value):
        return str(value if value not in [None, ""] else "Not provided")

    def wrap_text(text, font_name="Helvetica", font_size=10.2, max_width=None):
        max_width = max_width or usable_width
        words = safe_text(text).replace("\t", " ").split()
        lines = []
        current = ""
        for word in words:
            test = word if not current else current + " " + word
            if stringWidth(test, font_name, font_size) <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        return lines or [""]

    def draw_logo_or_brand(x, y_top, max_w=140, max_h=48, dark=False):
        logo_source = BytesIO(logo_bytes) if logo_bytes else logo_path
        if logo_source:
            try:
                img = ImageReader(logo_source)
                iw, ih = img.getSize()
                scale = min(max_w / iw, max_h / ih)
                draw_w, draw_h = iw * scale, ih * scale
                c.drawImage(img, x, y_top - draw_h, width=draw_w, height=draw_h, mask="auto", preserveAspectRatio=True)
                return draw_w, draw_h
            except Exception:
                pass
        c.setFont("Helvetica-Bold", 15)
        c.setFillColor(colors.white if dark else charcoal)
        c.drawString(x, y_top - 20, brand_practice)
        c.setFont("Helvetica", 7.5)
        c.setFillColor(colors.white if dark else mid_grey)
        c.drawString(x, y_top - 32, "Planning & Architectural Services")
        return max_w, 32

    def draw_footer():
        c.setStrokeColor(line_grey)
        c.setLineWidth(0.7)
        c.line(left_margin, 34, width - right_margin, 34)
        c.setFont("Helvetica", 8)
        c.setFillColor(mid_grey)
        footer_left = f"Prepared by {brand_practice} | Planning & Architectural Services"
        c.drawString(left_margin, 22, footer_left)
        c.drawRightString(width - right_margin, 22, f"Page {page_no}")

    def draw_header():
        c.setStrokeColor(line_grey)
        c.setLineWidth(0.7)
        c.line(left_margin, height - 36, width - right_margin, height - 36)
        draw_logo_or_brand(left_margin, height - 12, max_w=125, max_h=28)
        c.setFillColor(mid_grey)
        c.setFont("Helvetica", 8.2)
        c.drawRightString(width - right_margin, height - 20, f"{report_title} | Ref: {report_id}")

    def new_page():
        nonlocal y, page_no
        c.showPage()
        page_no += 1
        y = height - top_margin
        draw_header()
        draw_footer()

    def ensure_space(required_height):
        nonlocal y
        if y - required_height < bottom_margin:
            new_page()

    def draw_cover():
        nonlocal page_no
        page_no = 1
        # dark charcoal title band
        c.setFillColor(charcoal)
        c.rect(0, height - 230, width, 230, fill=1, stroke=0)
        # gold accent strip
        c.setFillColor(gold)
        c.rect(0, height - 235, width, 5, fill=1, stroke=0)

        draw_logo_or_brand(left_margin, height - 45, max_w=170, max_h=58, dark=True)
        c.setFillColor(colors.white)
        c.setFont("Helvetica", 10)
        c.drawRightString(width - right_margin, height - 55, SY_BRAND["product_name"])
        c.setFont("Helvetica", 8.5)
        c.drawRightString(width - right_margin, height - 70, f"Report ID: {report_id}")

        title = report_title
        if module_name == "Planning Review":
            title = "Planning Appraisal Report"
        elif module_name == "Building Regulations Review":
            title = "Building Regulations Review Report"

        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 24)
        title_lines = wrap_text(title, "Helvetica-Bold", 24, width - 2 * left_margin)
        yy = height - 125
        for line in title_lines[:2]:
            c.drawString(left_margin, yy, line)
            yy -= 29
        c.setFont("Helvetica", 11)
        c.setFillColor(colors.HexColor("#E8E2D5"))
        c.drawString(left_margin, yy - 4, "Prepared in a professional planning-consultant style using ArchLens AI")

        # Meta card
        card_x = left_margin
        card_y = height - 495
        card_w = usable_width
        card_h = 210
        c.setFillColor(colors.white)
        c.roundRect(card_x, card_y, card_w, card_h, 12, fill=1, stroke=0)
        c.setStrokeColor(line_grey)
        c.roundRect(card_x, card_y, card_w, card_h, 12, fill=0, stroke=1)
        c.setFillColor(light_gold)
        c.roundRect(card_x, card_y + card_h - 36, card_w, 36, 12, fill=1, stroke=0)
        c.setFillColor(charcoal)
        c.setFont("Helvetica-Bold", 11)
        c.drawString(card_x + 18, card_y + card_h - 23, "Project Information")

        meta = [
            ("Project Address", safe_text(address)),
            ("Client", safe_text(client)),
            ("Project Type", sections.get("PROJECT CLASSIFICATION", "Not detected").splitlines()[0] if sections.get("PROJECT CLASSIFICATION") else "Not detected"),
            ("Drawing Pack Reviewed", safe_text(file_name)),
            ("Date", safe_text(date)),
            ("Prepared By", brand_practice),
        ]
        yy = card_y + card_h - 62
        for label, value in meta:
            c.setFont("Helvetica-Bold", 8.8)
            c.setFillColor(mid_grey)
            c.drawString(card_x + 18, yy, label.upper())
            c.setFont("Helvetica", 9.4)
            c.setFillColor(charcoal)
            x_val = card_x + 155
            for i, line in enumerate(wrap_text(value, "Helvetica", 9.4, card_w - 175)[:2]):
                c.drawString(x_val, yy - (i * 11), line)
            yy -= 27

        # Summary card
        summary = sections.get("TOP SUMMARY", "") or sections.get("EXECUTIVE SUMMARY", "")
        c.setFillColor(soft_grey)
        c.roundRect(left_margin, 140, usable_width, 90, 10, fill=1, stroke=0)
        c.setFillColor(charcoal)
        c.setFont("Helvetica-Bold", 11)
        c.drawString(left_margin + 16, 206, "Executive Summary")
        c.setFont("Helvetica", 9.4)
        c.setFillColor(colors.HexColor("#444444"))
        yy = 190
        summary_lines = []
        for raw in summary.splitlines():
            raw = raw.strip(" -•")
            if raw:
                summary_lines.extend(wrap_text(raw, "Helvetica", 9.4, usable_width - 32))
        for line in summary_lines[:5]:
            c.drawString(left_margin + 16, yy, line)
            yy -= 12

        c.setFont("Helvetica", 8)
        c.setFillColor(mid_grey)
        c.drawCentredString(width / 2, 58, "This document is AI-assisted and should be reviewed by a competent professional before formal submission.")
        c.showPage()

    def draw_section_heading(number, title):
        nonlocal y
        ensure_space(48)
        c.setFillColor(light_gold)
        c.roundRect(left_margin, y - 27, usable_width, 30, 7, fill=1, stroke=0)
        c.setFillColor(gold)
        c.roundRect(left_margin, y - 27, 34, 30, 7, fill=1, stroke=0)
        c.setFillColor(charcoal)
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(left_margin + 17, y - 16, f"{number}")
        c.setFont("Helvetica-Bold", 12)
        c.drawString(left_margin + 46, y - 16, title)
        y -= 43

    def draw_paragraph(text, indent=0, bullet=False, font_size=10.2):
        nonlocal y
        max_w = usable_width - indent - (16 if bullet else 0)
        lines = wrap_text(text, "Helvetica", font_size, max_w)
        ensure_space(len(lines) * 14 + 12)
        c.setFont("Helvetica", font_size)
        c.setFillColor(colors.HexColor("#3F3F3F"))
        if bullet:
            c.setFillColor(gold)
            c.circle(left_margin + indent + 3, y - 3, 2.2, fill=1, stroke=0)
            c.setFillColor(colors.HexColor("#3F3F3F"))
            x = left_margin + indent + 14
        else:
            x = left_margin + indent
        for line in lines:
            c.drawString(x, y, line)
            y -= 14
        y -= 5

    def draw_key_value_section(content):
        nonlocal y
        rows = parse_key_value_lines(content)
        if not rows:
            draw_paragraph("Not detected")
            return
        for label, value in rows:
            ensure_space(32)
            if label:
                c.setFont("Helvetica-Bold", 9.2)
                c.setFillColor(mid_grey)
                c.drawString(left_margin, y, label.upper())
                y -= 13
                for line in wrap_text(value, "Helvetica", 10.2, usable_width - 10):
                    ensure_space(16)
                    c.setFont("Helvetica", 10.2)
                    c.setFillColor(colors.HexColor("#3F3F3F"))
                    c.drawString(left_margin + 10, y, line)
                    y -= 14
                y -= 5
            else:
                draw_paragraph(value)

    def draw_text_section(content):
        if not content or not content.strip():
            draw_paragraph("Not detected")
            return
        for raw in content.splitlines():
            line = raw.strip()
            if not line:
                continue
            if line.startswith("- ") or line.startswith("• "):
                draw_paragraph(line[2:].strip(), bullet=True)
            else:
                draw_paragraph(line)

    def draw_compliance_table(content):
        nonlocal y
        rows = parse_compliance_rows(content)
        if not rows:
            draw_text_section(content)
            return

        col_w = [45, 190, 105, usable_width - 45 - 190 - 105]
        x_positions = [left_margin]
        for w_col in col_w[:-1]:
            x_positions.append(x_positions[-1] + w_col)

        def table_header():
            nonlocal y
            ensure_space(31)
            c.setFillColor(charcoal)
            c.roundRect(left_margin, y - 22, usable_width, 24, 5, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 8.2)
            headers = ["PART", "APPROVED DOCUMENT", "STATUS", "COMMENTARY"]
            for x, h in zip(x_positions, headers):
                c.drawString(x + 7, y - 14, h)
            y -= 30

        table_header()
        for row in rows:
            doc_lines = wrap_text(row.get("title", ""), "Helvetica-Bold", 8.4, col_w[1] - 12)
            why_lines = wrap_text(row.get("why", ""), "Helvetica", 8.2, col_w[3] - 12)
            status_lines = wrap_text(row.get("status", ""), "Helvetica-Bold", 8, col_w[2] - 12)
            row_h = max(38, 14 + max(len(doc_lines), len(why_lines), len(status_lines), 1) * 10)
            ensure_space(row_h + 8)
            if y - row_h < bottom_margin:
                new_page()
                table_header()
            c.setFillColor(soft_grey)
            c.roundRect(left_margin, y - row_h + 4, usable_width, row_h, 5, fill=1, stroke=0)
            c.setStrokeColor(colors.white)
            for x in x_positions[1:]:
                c.line(x, y + 4, x, y - row_h + 4)
            c.setFillColor(charcoal)
            c.setFont("Helvetica-Bold", 9)
            c.drawString(x_positions[0] + 10, y - 13, row.get("part", ""))
            c.setFont("Helvetica-Bold", 8.4)
            yy = y - 12
            for line in doc_lines:
                c.drawString(x_positions[1] + 7, yy, line)
                yy -= 10
            status = row.get("status", "")
            status_upper = status.upper()
            fill = gold
            if "FAIL" in status_upper:
                fill = colors.HexColor("#D9A0A0")
            elif "PASS" in status_upper and "PARTLY" not in status_upper:
                fill = colors.HexColor("#B7D7B0")
            c.setFillColor(fill)
            c.roundRect(x_positions[2] + 7, y - 20, min(col_w[2] - 14, 92), 16, 4, fill=1, stroke=0)
            c.setFillColor(charcoal)
            c.setFont("Helvetica-Bold", 7.4)
            c.drawCentredString(x_positions[2] + 53, y - 15, status[:22].upper())
            c.setFillColor(colors.HexColor("#3F3F3F"))
            c.setFont("Helvetica", 8.2)
            yy = y - 12
            for line in why_lines:
                c.drawString(x_positions[3] + 7, yy, line)
                yy -= 10
            y -= row_h + 7

    draw_cover()
    page_no = 2
    y = height - top_margin
    draw_header()
    draw_footer()

    for idx, (key, title) in enumerate(section_order, start=1):
        content = sections.get(key, "Not detected")
        draw_section_heading(idx, title)
        if module_name == "Building Regulations Review" and key == "COMPLIANCE STATUS BY APPROVED DOCUMENT":
            draw_compliance_table(content)
        elif key in special_key_value_sections:
            draw_key_value_section(content)
        else:
            draw_text_section(content)
        y -= 6

    # Final professional disclaimer page note when space allows
    ensure_space(48)
    c.setFillColor(soft_grey)
    c.roundRect(left_margin, y - 40, usable_width, 42, 7, fill=1, stroke=0)
    c.setFillColor(mid_grey)
    c.setFont("Helvetica-Oblique", 8.4)
    yy = y - 14
    for line in wrap_text(disclaimer_text, "Helvetica-Oblique", 8.4, usable_width - 24)[:2]:
        c.drawString(left_margin + 12, yy, line)
        yy -= 11

    c.save()
    buffer.seek(0)
    return buffer


def build_word_report(file_name, address, client, date, practice_name, report_id, sections, module_name):
    config = MODULE_CONFIG[module_name]
    section_order = config["section_order"]
    special_key_value_sections = config["special_key_value_sections"]
    disclaimer_text = config["disclaimer"]
    report_title = config["title"]

    doc = Document()

    section = doc.sections[0]
    add_word_header(section, practice_name or "ArchLens AI", report_title)
    add_word_footer(section, practice_name or "ArchLens AI", report_title)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(practice_name or "ArchLens AI")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(report_title)
    run.bold = True
    run.font.size = Pt(16)

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(disclaimer_text)
    run.font.size = Pt(9)
    run.italic = True

    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Project Address: {address}\n").bold = True
    info.add_run(f"Client: {client}\n")
    info.add_run(f"Drawing Pack Reviewed: {file_name}\n")
    info.add_run(f"Date: {date}\n")
    info.add_run(f"Report ID: {report_id}\n")
    info.add_run(f"Prepared by: {practice_name or 'ArchLens AI'}")

    doc.add_paragraph("")
    summary_heading = doc.add_paragraph()
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary_heading.add_run("Client Summary")
    run.bold = True
    run.font.size = Pt(13)

    for label, value in parse_key_value_lines(sections.get("TOP SUMMARY", "")):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if label:
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)
        else:
            p.add_run(value)

    doc.add_page_break()

    for section in doc.sections:
        add_word_header(section, practice_name or "ArchLens AI", report_title)
        add_word_footer(section, practice_name or "ArchLens AI", report_title)

    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_col_widths(meta_table, [2.0, 4.8])
    meta_rows = [
        ("Project Address", str(address)),
        ("Client", str(client)),
        ("Drawing Pack Reviewed", str(file_name)),
        ("Date", str(date)),
        ("Report ID", str(report_id)),
        ("Prepared by", str(practice_name or "ArchLens AI")),
    ]
    for i, (label, value) in enumerate(meta_rows):
        left_cell = meta_table.rows[i].cells[0]
        right_cell = meta_table.rows[i].cells[1]
        set_cell_text(left_cell, label, bold=True, font_size=10)
        set_cell_text(right_cell, value, font_size=10)
        set_cell_shading(left_cell, "EAEFF7")
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph("")

    for key, title in section_order:
        doc.add_paragraph("")
        heading = doc.add_paragraph()
        heading_run = heading.add_run(title)
        heading_run.bold = True
        heading_run.font.size = Pt(13)
        content = sections.get(key, "Not detected")

        if key in special_key_value_sections:
            rows = parse_key_value_lines(content)
            for label, value in rows:
                if label:
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)
                else:
                    doc.add_paragraph(value)

        elif module_name == "Building Regulations Review" and key == "COMPLIANCE STATUS BY APPROVED DOCUMENT":
            rows = parse_compliance_rows(content)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            set_table_col_widths(table, [0.8, 2.3, 1.6, 3.3])

            headers = table.rows[0].cells
            set_cell_text(headers[0], "Part", bold=True, font_size=9)
            set_cell_text(headers[1], "Approved Document", bold=True, font_size=9)
            set_cell_text(headers[2], "Status", bold=True, font_size=9)
            set_cell_text(headers[3], "Why", bold=True, font_size=9)
            for cell in headers:
                set_cell_shading(cell, "D9E2F3")

            for row in rows:
                cells = table.add_row().cells
                set_cell_text(cells[0], row["part"], bold=True, font_size=9)
                set_cell_text(cells[1], row["title"], bold=True, font_size=9)
                set_cell_text(cells[2], row["status"], bold=True, font_size=9)
                set_cell_text(cells[3], row["why"], font_size=9)
                status_upper = row["status"].upper()
                if "PASS" in status_upper and "PARTLY" not in status_upper:
                    set_cell_shading(cells[2], "C6EFCE")
                elif "FAIL" in status_upper:
                    set_cell_shading(cells[2], "FFC7CE")
                elif "PARTLY" in status_upper or "REVIEW REQUIRED" in status_upper:
                    set_cell_shading(cells[2], "FCE4D6")
        else:
            for line in content.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("- "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                else:
                    doc.add_paragraph(stripped)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def extract_summary_value(sections: Dict[str, str], module_name: str):
    top_summary_rows = {k.upper(): v for k, v in parse_key_value_lines(sections.get("TOP SUMMARY", "")) if k}
    if module_name == "Planning Review":
        authority_value = top_summary_rows.get("LOCAL AUTHORITY", "Unknown")
        if authority_value == "Unknown":
            for line in sections.get("TOP SUMMARY", "").splitlines():
                stripped = line.strip()
                if stripped and ":" not in stripped:
                    authority_value = stripped
                    break
        return (
            "Not shown",
            top_summary_rows.get("APPLICATION TYPE", top_summary_rows.get("LIKELY ROUTE", "Unknown")),
            authority_value,
        )
    return (
        top_summary_rows.get("OVERALL RISK RATING", "Unknown"),
        top_summary_rows.get("SUBMISSION STATUS", "Unknown"),
        top_summary_rows.get("REVIEW CONFIDENCE", "Unknown"),
    )


def render_kpi_cards(sections: Dict[str, str], report_id: str, module_name: str):
    v1, v2, v3 = extract_summary_value(sections, module_name)
    label_2 = "Likely Route" if module_name == "Planning Review" else "Submission Status"
    label_3 = "Local Authority" if module_name == "Planning Review" else "Review Confidence"

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Report ID", report_id)
    if module_name == "Planning Review":
        c2.metric("Planning Position", v1)
        c3.metric(label_2, v2)
        c4.metric(label_3, v3)
    else:
        c2.metric("Risk Rating", v1)
        c3.metric(label_2, v2)
        c4.metric(label_3, v3)


def extract_summary_values(sections: Dict[str, str], module_name: str):
    top_summary_rows = {k.upper(): v for k, v in parse_key_value_lines(sections.get("TOP SUMMARY", "")) if k}
    if module_name == "Planning Review":
        authority_value = top_summary_rows.get("LOCAL AUTHORITY", "Unknown")
        if authority_value == "Unknown":
            for line in sections.get("TOP SUMMARY", "").splitlines():
                stripped = line.strip()
                if stripped and ":" not in stripped:
                    authority_value = stripped
                    break
        return {
            "risk": "Not shown",
            "route": top_summary_rows.get("APPLICATION TYPE", top_summary_rows.get("LIKELY ROUTE", "Unknown")),
            "authority": authority_value,
            "probability": top_summary_rows.get("PLANNING ROUTE CONFIDENCE SCORE", "Not shown"),
        }
    return {
        "risk": top_summary_rows.get("OVERALL RISK RATING", "Unknown"),
        "route": top_summary_rows.get("SUBMISSION STATUS", "Unknown"),
        "authority": top_summary_rows.get("REVIEW CONFIDENCE", "Unknown"),
        "probability": "N/A",
    }


def detect_local_authority_for_display(project_address: str, proposal_summary: str, uploaded_files=None) -> str:
    combined_text = proposal_summary or ""
    if uploaded_files:
        names = " ".join(f.name for f in uploaded_files if getattr(f, "name", None))
        combined_text = f"{combined_text}\n{names}"
    return pdf_summary.detect_local_authority(project_address or "", combined_text or "")

def render_at_a_glance(sections: Dict[str, str], report_id: str, module_name: str):
    config = MODULE_CONFIG[module_name]
    readiness_key = config["readiness_key"]
    middle_key = "PROJECT DETAILS" if module_name == "Building Regulations Review" else "SITE AND PROPOSAL OVERVIEW"
    middle_title = "Project Details" if module_name == "Building Regulations Review" else "Site and Proposal Overview"

    render_kpi_cards(sections, report_id, module_name)
    st.markdown("")

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown('<div class="sy-mini-card"><div class="sy-kpi">Project Classification</div></div>', unsafe_allow_html=True)
        st.markdown(sections.get("PROJECT CLASSIFICATION", "Not detected"))
    with col2:
        st.markdown(f'<div class="sy-mini-card"><div class="sy-kpi">{middle_title}</div></div>', unsafe_allow_html=True)
        st.markdown(sections.get(middle_key, "Not detected"))
    with col3:
        st.markdown('<div class="sy-mini-card"><div class="sy-kpi">Submission Readiness</div></div>', unsafe_allow_html=True)
        st.markdown(sections.get(readiness_key, "Not detected"))

    st.markdown("")
    if module_name == "Planning Review":
        readiness_text = sections.get("SUBMISSION READINESS", "")
        if "READY TO SUBMIT" in readiness_text.upper():
            st.success("Submission position: Ready to submit")
        elif "LIKELY READY" in readiness_text.upper():
            st.warning("Submission position: Likely ready with minor amendments")
        else:
            st.info("Submission position: Further information required")
    else:
        top_summary_rows = {k.upper(): v for k, v in parse_key_value_lines(sections.get("TOP SUMMARY", "")) if k}
        risk_summary = top_summary_rows.get("OVERALL RISK RATING") or sections.get("TOP SUMMARY", "")
        summary_hint = top_summary_rows.get("REVIEW CONFIDENCE", "Unknown")
        if "HIGH" in str(risk_summary).upper():
            st.error(f"High risk detected | Summary: {summary_hint}")
        elif "MEDIUM" in str(risk_summary).upper():
            st.warning(f"Moderate risk detected | Summary: {summary_hint}")
        else:
            st.success(f"Lower risk indicated | Summary: {summary_hint}")


def render_section_content(content: str, is_key_value: bool):
    if is_key_value:
        rows = parse_key_value_lines(content)
        for label, value in rows:
            if label:
                st.markdown(f"**{label}:** {value}")
            else:
                st.markdown(value)
    else:
        st.markdown(content)


def render_sections(sections: Dict[str, str], report_text: str, module_name: str):
    config = MODULE_CONFIG[module_name]
    for key, title in config["section_order"]:
        content = sections.get(key, "Not detected")
        with st.expander(title, expanded=key in {"TOP SUMMARY", config["readiness_key"]}):
            render_section_content(content, key in config["special_key_value_sections"])
    with st.expander("Show full AI report"):
        st.text(report_text)


def build_simple_word_doc(title: str, body_text: str) -> BytesIO:
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("")
    for line in body_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


st.set_page_config(page_title="ArchLens AI", layout="wide")
inject_custom_css()

if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False
if "auth_plan" not in st.session_state:
    st.session_state["auth_plan"] = "starter"
if "auth_user_name" not in st.session_state:
    st.session_state["auth_user_name"] = ""

token_plan, token_user_name, has_valid_token = get_verified_plan_and_user()

if has_valid_token:
    st.session_state["authenticated"] = True
    st.session_state["auth_plan"] = token_plan
    st.session_state["auth_user_name"] = token_user_name

if not st.session_state.get("authenticated", False):
    st.markdown(
        """
        <div class="sy-hero" style="max-width:900px;margin:3rem auto 1rem auto;">
            <div class="sy-hero-copy">
                <h1>ArchLens AI</h1>
                <div class="sy-muted" style="max-width:760px;">
                    Access is managed through your SY Design Studio member account.
                    Please launch ArchLens from your member area to verify your subscription and open the correct plan.
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    col1, col2 = st.columns(2)
    with col1:
        st.link_button("Login to Member Area", WEBSITE_LAUNCH_URL, use_container_width=True)
    with col2:
        st.link_button("View Plans", WEBSITE_PRICING_URL, use_container_width=True)
    st.stop()

current_plan = st.session_state.get("auth_plan", "starter")
current_user_name = st.session_state.get("auth_user_name", "")
if current_user_name:
    sync_credit_balance_from_api(current_user_name)
allowed_review_modules = get_allowed_review_modules(current_plan)

# -------------------------------
# ArchGuard/ArchLens dashboard UI
# -------------------------------
WIZARD_DEFAULTS = {
    "app_page": "Projects",
    "project_step": 1,
    "wizard_review_module": allowed_review_modules[0],
    "wizard_review_mode": "Architect / Professional",
    "wizard_project_name": "",
    "wizard_project_address": "",
    "wizard_client_name": "",
    "wizard_project_types": [],
    "wizard_property_type": "Not stated",
    "wizard_proposal_summary": "",
    "wizard_rear_depth": 6.0,
    "wizard_rear_height": 4.0,
    "wizard_uploaded_files": [],
    "wizard_scope_items": [],
    "wizard_review_focus": "",
    "wizard_accuracy_answers": {},
}
for k, v in WIZARD_DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

if st.session_state["wizard_review_module"] not in allowed_review_modules:
    st.session_state["wizard_review_module"] = allowed_review_modules[0]

def get_brand_logo_bytes_for_ui():
    logo_bytes = st.session_state.get("brand_logo_bytes")
    if logo_bytes:
        return logo_bytes
    logo_path = get_brand_logo_path()
    if logo_path and os.path.exists(logo_path):
        with open(logo_path, "rb") as f:
            return f.read()
    return None

def app_logo_data_uri():
    logo_bytes = get_brand_logo_bytes_for_ui()
    if not logo_bytes:
        return ""
    return "data:image/png;base64," + base64.b64encode(logo_bytes).decode("utf-8")

def checkbox_grid(label, options, state_key, columns=2, help_text=None):
    st.markdown(f"**{label}**")
    if help_text:
        st.caption(help_text)
    current = set(st.session_state.get(state_key, []))
    selected = []
    cols = st.columns(columns)
    for i, option in enumerate(options):
        with cols[i % columns]:
            checked = st.checkbox(option, value=option in current, key=f"{state_key}_{i}")
            if checked:
                selected.append(option)
    st.session_state[state_key] = selected
    return selected


def single_choice_cards(label, options, state_key, columns=2, help_text=None):
    st.markdown(f"**{label}**")
    if help_text:
        st.caption(help_text)
    current = st.session_state.get(state_key, options[0] if options else "")
    cols = st.columns(columns)
    for i, option in enumerate(options):
        with cols[i % columns]:
            if st.checkbox(option, value=(current == option), key=f"{state_key}_{i}"):
                current = option
    st.session_state[state_key] = current
    return current


def get_required_accuracy_answers(project_types):
    group = get_accuracy_question_group(project_types)
    answers = {"pd_question_family": group}
    st.markdown("**Planning / PD accuracy questions**")
    st.caption("These answers are part of the project intake and will be used with the uploaded drawings. Where drawing dimensions differ, the uploaded plans take priority.")
    if group in {"class_a", "class_b", "class_d"}:
        answers["site_constraints"] = checkbox_grid(
            "Site constraints",
            ["None", "Conservation area / Article 2(3) land", "Article 4 direction", "Listed building"],
            "wizard_acc_site_constraints",
            columns=2,
        )
    if group == "class_b":
        answers["front_roof_plane_highway"] = st.radio("Does any part project from the front roof slope facing the highway?", ["Not sure", "No", "Yes"], horizontal=True, key="wizard_acc_b_front")
        answers["roof_volume_band"] = st.radio("Added roof volume", ["Not sure", "Within normal limit", "Over limit"], horizontal=True, key="wizard_acc_b_volume")
        answers["above_existing_roof_height"] = st.radio("Does it rise above the highest part of the existing roof?", ["Not sure", "No", "Yes"], horizontal=True, key="wizard_acc_b_highest")
        answers["eaves_setback_0_2m"] = st.radio("Is the enlargement set back at least 200mm from the original eaves?", ["Not sure", "Yes", "No"], horizontal=True, key="wizard_acc_b_eaves")
        answers["side_windows_obscure_glazed"] = st.radio("Any side windows obscure glazed and non-opening below 1.7m?", ["Not applicable", "Yes", "No", "Not sure"], horizontal=True, key="wizard_acc_b_windows")
        answers["materials_similar"] = st.radio("Will external materials be similar in appearance to the existing house?", ["Not sure", "Yes", "No"], horizontal=True, key="wizard_acc_b_materials")
    elif group == "class_a":
        answers["forward_of_principal_elevation"] = st.radio("Does any part project beyond the principal elevation or a side elevation facing a highway?", ["Not sure", "No", "Yes"], horizontal=True, key="wizard_acc_a_principal")
        answers["existing_rear_extensions"] = st.radio("Have there been previous rear extensions added to the original house?", ["Not sure", "No", "Yes"], horizontal=True, key="wizard_acc_a_existing")
        answers["within_2m_of_boundary"] = st.radio("Is any part within 2m of a boundary?", ["Not sure", "No", "Yes"], horizontal=True, key="wizard_acc_a_boundary")
        answers["eaves_height_within_2m"] = st.radio("If within 2m of a boundary, are the eaves 3.0m or lower?", ["Not applicable", "Yes", "No", "Not sure"], horizontal=True, key="wizard_acc_a_eaves")
        answers["side_extension_width"] = st.radio("For a side extension, is it more than half the width of the original house?", ["Not applicable", "No", "Yes", "Not sure"], horizontal=True, key="wizard_acc_a_width")
        answers["materials_similar"] = st.radio("Will external materials be similar in appearance to the existing house?", ["Not sure", "Yes", "No"], horizontal=True, key="wizard_acc_a_materials")
    elif group == "class_d":
        answers["porch_ground_area_band"] = st.radio("Is the porch ground area 3m² or less?", ["Not sure", "Yes", "No"], horizontal=True, key="wizard_acc_d_area")
        answers["porch_height_band"] = st.radio("Is the porch 3.0m high or lower?", ["Not sure", "Yes", "No"], horizontal=True, key="wizard_acc_d_height")
        answers["porch_within_2m_highway"] = st.radio("Is any part within 2m of a boundary with a highway?", ["Not sure", "No", "Yes"], horizontal=True, key="wizard_acc_d_highway")
    else:
        st.info("No specific PD question set has been triggered for this project type. Add any constraints or review focus below.")
    return answers


def render_left_navigation():
    logo_uri = app_logo_data_uri()
    theme = st.session_state.get("app_theme", "Dark")
    light_mode = str(theme).lower().startswith("light")
    logo_box_bg = "transparent" if light_mode else "#061225"
    logo_padding = "0px" if light_mode else "8px"
    logo_subtitle_colour = "#6B7280" if light_mode else "#9FB2D8"
    with st.sidebar:
        if logo_uri:
            st.markdown(
                f'''
                <div style="display:flex;align-items:center;gap:1rem;margin:0.9rem 0 1.45rem 0;">
                    <img src="{logo_uri}" style="width:112px;height:112px;object-fit:contain;border-radius:18px;background:{logo_box_bg};padding:{logo_padding};" />
                    <div>
                        <div style="font-weight:850;font-size:1.28rem;line-height:1.15;">ArchLens AI</div>
                        <div style="font-size:0.86rem;color:{logo_subtitle_colour};margin-top:0.25rem;">SY Design Studio</div>
                    </div>
                </div>
                ''',
                unsafe_allow_html=True,
            )
        else:
            st.markdown("### ArchLens AI")
        st.caption(f"Plan: {PLAN_LABELS.get(current_plan, 'Solo')}")
        st.caption(f"Credits: {get_credit_balance()}")
        if current_user_name:
            st.caption(f"User: {current_user_name}")
        page = st.radio(
            "Navigation",
            ["Dashboard", "Projects", "Reports", "Settings"],
            index=["Dashboard", "Projects", "Reports", "Settings"].index(st.session_state.get("app_page", "Projects")),
            label_visibility="collapsed",
        )
        st.session_state["app_page"] = page
        st.markdown("---")
        st.link_button("Return to SY Design Studio", WEBSITE_HOME_URL, use_container_width=True)
        st.link_button("Buy Credits", ARCHLENS_BUY_CREDITS_URL, use_container_width=True)
    return page

def intake_items():
    review_module = st.session_state.get("wizard_review_module")
    return [
        ("Module selected", bool(review_module)),
        ("Project named", bool(st.session_state.get("wizard_project_name") or st.session_state.get("wizard_project_address"))),
        ("Project type selected", bool(st.session_state.get("wizard_project_types"))),
        ("Property details captured", review_module == "Building Regulations Review" or st.session_state.get("wizard_property_type") != "Not stated"),
        ("Site and council added", bool(st.session_state.get("wizard_project_address"))),
        ("Scope selected", bool(st.session_state.get("wizard_proposal_summary") or st.session_state.get("wizard_project_types"))),
        ("Files uploaded", bool(st.session_state.get("wizard_uploaded_files"))),
    ]

def render_intake_panel():
    items = intake_items()
    complete = sum(1 for _, done in items if done)
    total = len(items)
    project_types = st.session_state.get("wizard_project_types", [])
    project_address = st.session_state.get("wizard_project_address", "")
    proposal_summary = st.session_state.get("wizard_proposal_summary", "")
    local_authority = detect_local_authority_for_display(project_address, proposal_summary)
    st.markdown('<div class="sy-sidepanel">', unsafe_allow_html=True)
    st.markdown('<div class="sy-panel-title">Intake Readiness</div>', unsafe_allow_html=True)
    st.markdown(f"**{complete} of {total} key items completed**")
    for label, done in items:
        icon = "✅" if done else "•"
        st.markdown(f"{icon} {label}")
    st.progress(complete / max(total, 1))
    st.markdown("---")
    st.markdown('<div class="sy-panel-title">Live Summary</div>', unsafe_allow_html=True)
    rows = [
        ("Project", st.session_state.get("wizard_project_name") or "Not named"),
        ("Module", st.session_state.get("wizard_review_module", "Not selected")),
        ("Project type", ", ".join(project_types) if project_types else "Not selected"),
        ("Property", st.session_state.get("wizard_property_type", "Not stated")),
        ("Site", project_address or "Not added"),
        ("Council", local_authority or "Not detected"),
    ]
    for label, value in rows:
        st.markdown(f'<div class="sy-data-row"><span>{label}</span><strong>{value}</strong></div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

def step_header(step_no, title, subtitle):
    st.markdown(
        f'''
        <div class="sy-subtle-card">
            <div class="sy-section-label">Step {step_no}</div>
            <h2 style="margin:0 0 0.35rem 0;">{title}</h2>
            <div class="sy-muted">{subtitle}</div>
        </div>
        ''',
        unsafe_allow_html=True,
    )

def wizard_buttons(max_step=7):
    c1, c2, c3 = st.columns([0.35, 1, 0.35])
    with c1:
        if st.session_state.project_step > 1:
            if st.button("Back", use_container_width=True):
                st.session_state.project_step -= 1
                st.rerun()
    with c3:
        if st.session_state.project_step < max_step:
            if st.button("Continue", use_container_width=True):
                st.session_state.project_step += 1
                st.rerun()

def run_archlens_analysis(uploaded_files):
    review_module = st.session_state.get("wizard_review_module")
    review_mode = st.session_state.get("wizard_review_mode")
    project_types = st.session_state.get("wizard_project_types", [])
    property_type = st.session_state.get("wizard_property_type", "Not stated")
    proposal_summary = st.session_state.get("wizard_proposal_summary", "")
    project_address = st.session_state.get("wizard_project_address", "")
    client_name = st.session_state.get("wizard_client_name", "")
    rear_extension_depth_m = float(st.session_state.get("wizard_rear_depth", 0) or 0)
    rear_extension_height_m = float(st.session_state.get("wizard_rear_height", 0) or 0)
    local_authority = detect_local_authority_for_display(project_address, proposal_summary, uploaded_files)
    accuracy_answers = st.session_state.get("wizard_accuracy_answers", {}) or {}
    scope_items = st.session_state.get("wizard_scope_items", []) or []
    review_focus = st.session_state.get("wizard_review_focus") or ""
    rule_engine_summary = ""
    drawing_priority_instruction = (
        "Important instruction: cross-check all user-entered project type, scope items and measurements against the uploaded drawing PDF. "
        "If the uploaded plans show different dimensions or scope, the uploaded plans take priority. "
        "Only state measurements and conclusions that are supported by the drawings, policy, guidance or Building Regulations. "
        "If the user gives a specific review focus, concentrate the report on that issue and avoid unsupported assumptions."
    )
    config = MODULE_CONFIG[review_module]

    # Token system: module access is no longer blocked by Solo/Studio here.
    # Keep file-size/page limits below to protect AI/API costs.
    if not uploaded_files:
        st.error("Please upload at least one PDF drawing pack before running the review.")
        st.stop()

    total_uploaded_mb = sum(f.size for f in uploaded_files) / (1024 * 1024)
    if total_uploaded_mb > 20:
        st.error("Drawing pack too large. Please keep the total upload size to 20MB or less, or split the pack into smaller PDFs.")
        st.stop()

    progress_bar = st.progress(0)
    status_text = st.empty()
    temp_pdf_path = None
    file = uploaded_files[-1]

    for f in uploaded_files:
        if f.size > MAX_FILE_SIZE_MB * 1024 * 1024:
            st.error(f"PDF too large. Maximum file size is {MAX_FILE_SIZE_MB}MB.")
            st.stop()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(f.getbuffer())
            temp_pdf_path = tmp_file.name

    try:
        page_count = get_pdf_page_count(temp_pdf_path)
        if page_count > MAX_PAGE_COUNT:
            st.error(f"PDF has {page_count} pages. Maximum allowed is {MAX_PAGE_COUNT} pages.")
            os.remove(temp_pdf_path)
            st.stop()
        if page_count > 12:
            st.warning("Large drawing pack detected. The live app will analyse the first 12 pages only to keep processing stable.")

        smooth_progress(progress_bar, status_text, 10, 25, "Preparing drawing analysis...", 0.6)

        def update_analysis_progress(current_batch, total_batches):
            start_pct = 25
            end_pct = 85
            progress = start_pct + int((current_batch / max(1, total_batches)) * (end_pct - start_pct))
            progress_bar.progress(progress)
            status_text.text(f"Step 3 of 4 — Analyzing drawing pages... Batch {current_batch} of {total_batches} ({progress}%)")

        try:
            if review_module == "Building Regulations Review":
                context = (
                    "Project type: " + (", ".join(project_types) or "Not stated") +
                    "\nSelected scope items: " + (", ".join(scope_items) or "Not stated") +
                    "\nProposal summary: " + (proposal_summary or "Not stated") +
                    "\nSpecific review focus / notes: " + (review_focus or "Not stated") +
                    "\n" + drawing_priority_instruction
                )
                report = pdf_summary.analyze_pdf(
                    temp_pdf_path,
                    client_project_type=context,
                    review_mode=review_mode,
                    progress_callback=update_analysis_progress,
                )
            else:
                smooth_progress(progress_bar, status_text, 25, 40, "Reading drawings and extracting planning data...", 0.8)
                proposal_summary_for_ai = proposal_summary
                if "Ground Floor Rear Extension" in project_types:
                    depth_txt = f"approx. user-entered {rear_extension_depth_m:.1f}m depth from rear wall" if rear_extension_depth_m else ""
                    height_txt = f"approx. user-entered {rear_extension_height_m:.1f}m overall height" if rear_extension_height_m else ""
                    extra_bits = ", ".join([x for x in [depth_txt, height_txt] if x])
                    if extra_bits:
                        proposal_summary_for_ai = (proposal_summary_for_ai.strip() + " | " + extra_bits + " | Drawing dimensions take priority if different.").strip(" |")
                if scope_items:
                    proposal_summary_for_ai = (proposal_summary_for_ai.strip() + " | Scope noted: " + ", ".join(scope_items)).strip(" |")
                if review_focus:
                    proposal_summary_for_ai = (proposal_summary_for_ai.strip() + " | Specific review focus / notes: " + review_focus).strip(" |")
                # Do not pass internal drawing-priority instructions into report wording.
                # The report prompt already includes this rule internally.
                proposal_summary_for_ai = proposal_summary_for_ai.strip(" |")
                pd_context = build_pd_context(project_types, property_type, rear_extension_depth_m, rear_extension_height_m, accuracy_answers)
                accuracy_context = build_accuracy_context(accuracy_answers)
                if accuracy_context:
                    proposal_summary_for_ai = (proposal_summary_for_ai.strip() + " | PD answers: " + accuracy_context).strip(" |")

                # Run deterministic householder PD rule checks before the AI narrative.
                # AI should explain these results, not replace them.
                try:
                    rule_facts = planning_rules.facts_from_app_context(
                        project_types=project_types,
                        property_type=property_type,
                        proposal_summary=proposal_summary_for_ai,
                        pd_context=pd_context,
                        scope_items=scope_items,
                    )
                    rule_result = planning_rules.run_householder_pd_rules(rule_facts)
                    rule_engine_summary = planning_rules.format_rule_result_for_prompt(rule_result)
                except Exception as rule_error:
                    rule_engine_summary = f"DETERMINISTIC RULE ENGINE RESULT: NEEDS CONFIRMATION\nSUMMARY: Rule engine could not complete: {rule_error}"

                report = pdf_summary.analyze_planning_pdf(
                    temp_pdf_path,
                    client_project_types=project_types,
                    property_type=property_type,
                    proposal_summary=proposal_summary_for_ai,
                    project_address=project_address,
                    local_authority=local_authority,
                    review_mode=review_mode,
                    pd_context=pd_context,
                    scope_items=scope_items,
                    rule_engine_summary=rule_engine_summary,
                )
                smooth_progress(progress_bar, status_text, 40, 85, "Analyzing planning route and risks...", 0.8)
        except Exception as e:
            msg = str(e).lower()
            if "insufficient_quota" in msg or "quota" in msg:
                st.error("OpenAI API quota exceeded. Please add credits in your OpenAI billing dashboard.")
            elif "rate limit" in msg or "429" in msg:
                st.error("The AI analysis service is temporarily rate-limited. Please try again shortly.")
            else:
                st.error(f"Could not analyze this PDF: {e}")
            st.stop()

        valid, missing = validate_report_headings(report, config["required_headings"])
        if not valid:
            st.error(f"AI report validation failed. Missing headings: {', '.join(missing)}")
            st.stop()

        sections = parse_report_sections(report, config["required_headings"])
        ai_confidence = calculate_ai_confidence(review_module, sections, rule_engine_summary)
        extracted_report_address = extract_address_from_report(report, "Not provided")
        clean_project_address = clean_input_value(project_address, extracted_report_address)
        clean_client_name = clean_input_value(client_name, "Not provided")
        clean_practice_name = "SY Design Studio"
        report_id = str(uuid.uuid4())[:8].upper()

        smooth_progress(progress_bar, status_text, 85, 95, "Preparing report files...", 0.6)
        word_file = build_word_report(file.name, clean_project_address, clean_client_name, time.strftime("%Y-%m-%d"), clean_practice_name, report_id, sections, review_module)
        pdf_file = build_pdf_report(file.name, clean_project_address, clean_client_name, time.strftime("%Y-%m-%d"), clean_practice_name, report_id, sections, review_module)

        st.session_state.report = report
        st.session_state.sections = sections
        st.session_state.word_file = word_file
        st.session_state.pdf_file = pdf_file
        st.session_state.last_filename = file.name
        st.session_state.report_id = report_id
        st.session_state.active_module = review_module
        st.session_state.ai_confidence = ai_confidence
        st.session_state.rule_engine_summary = rule_engine_summary
        st.session_state["planning_statement_text"] = None
        st.session_state["planning_statement_file"] = None
        if current_plan == "starter":
            st.session_state["starter_review_count"] = st.session_state.get("starter_review_count", 0) + 1
        add_saved_project({
            "report_id": report_id,
            "project_address": clean_project_address,
            "client_name": clean_client_name,
            "module": review_module,
            "project_types": ", ".join(project_types) if project_types else "Not stated",
            "property_type": property_type if review_module == "Planning Review" else "Not stated",
            "filename": file.name,
            "date": time.strftime("%Y-%m-%d"),
            "plan": PLAN_LABELS.get(current_plan, "Solo"),
            "local_authority": local_authority,
            "pdf_bytes": pdf_file.getvalue(),
            "word_bytes": word_file.getvalue(),
            "ai_confidence": ai_confidence,
            "pdf_unlocked": False,
            "word_unlocked": False,
            "credits_used": 0,
        })
        smooth_progress(progress_bar, status_text, 95, 100, "Finalising report...", 0.4)
        status_text.text("Analysis complete. 100%")
        progress_bar.progress(100)
        st.success("Report created successfully. Open Reports or stay on this page to download it.")
    finally:
        if temp_pdf_path:
            try:
                os.remove(temp_pdf_path)
            except Exception:
                pass
        gc.collect()

def render_report_download_panel(module_name=None, show_sections=True):
    if not st.session_state.get("sections"):
        st.info("No report generated yet.")
        return
    sections = st.session_state.sections
    report = st.session_state.report
    report_id = st.session_state.report_id or "N/A"
    module_name = module_name or st.session_state.get("active_module", "Planning Review")
    st.markdown('<div class="sy-subtle-card"><div class="sy-section-label">Review Output</div><h3 style="margin:0 0 0.35rem 0;">Latest Professional Report</h3><div class="sy-muted">Download the branded PDF or review the AI report sections.</div></div>', unsafe_allow_html=True)
    render_ai_confidence_card(st.session_state.get("ai_confidence"))
    render_at_a_glance(sections, report_id, module_name)
    base_filename = (st.session_state.last_filename or "drawing_pack").rsplit(".", 1)[0]
    suffix = "Planning" if module_name == "Planning Review" else "BuildingRegs"
    c1, c2 = st.columns(2)
    pdf_cost = get_export_credit_cost(module_name, "pdf")
    word_cost = get_export_credit_cost(module_name, "word")
    with c1:
        if is_report_unlocked(report_id, "pdf"):
            st.download_button("Download Branded PDF Report", st.session_state.pdf_file, file_name=f"{base_filename}_{suffix}_AI_Review_Report.pdf", mime="application/pdf", use_container_width=True)
        else:
            if st.button(f"Unlock PDF Report — {pdf_cost} credits", use_container_width=True):
                ok, msg = unlock_report_export(report_id, module_name, "pdf")
                if ok:
                    st.success(msg)
                    st.rerun()
                else:
                    st.error(msg)
    with c2:
        if is_report_unlocked(report_id, "word"):
            st.download_button("Download Word Report", st.session_state.word_file, file_name=f"{base_filename}_{suffix}_AI_Review_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        else:
            if st.button(f"Unlock Word Report — {word_cost} credit", use_container_width=True):
                ok, msg = unlock_report_export(report_id, module_name, "word")
                if ok:
                    st.success(msg)
                    st.rerun()
                else:
                    st.error(msg)
    if show_sections:
        render_sections(sections, report, module_name)

page = render_left_navigation()

# Main shell header
st.markdown(
    f'''
    <div class="sy-topbar">
        <div>
            <div class="sy-topbar-title">Architect AI Workspace</div>
            <div class="sy-topbar-meta">AI-powered planning and building regulations intelligence for UK projects</div>
        </div>
        <div class="sy-topbar-meta">Credits: {get_credit_balance()} | Plan: {PLAN_LABELS.get(current_plan, "Solo")}{(" | User: " + current_user_name) if current_user_name else ""}</div>
    </div>
    ''',
    unsafe_allow_html=True,
)

if page == "Dashboard":
    st.markdown('<div class="sy-hero"><div class="sy-hero-copy"><h1>Dashboard</h1><div class="sy-muted">Monitor your project intake, recent reports and subscription access.</div></div></div>', unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Saved projects", len(st.session_state.get("saved_projects", [])))
    c2.metric("Reports generated", len(st.session_state.get("saved_projects", [])))
    c3.metric("Credits", get_credit_balance())
    c4.metric("Current plan", PLAN_LABELS.get(current_plan, "Solo"))
    st.markdown("### Recent projects")
    saved_projects = st.session_state.get("saved_projects", [])
    if saved_projects:
        for item in saved_projects[:6]:
            st.markdown(f"**{item.get('project_address', 'Not provided')}**")
            conf = item.get("ai_confidence") or {}
            status_text = f" • Status: {conf.get('label')}" if conf.get("label") else ""
            st.caption(f"{item.get('module', '')} • {item.get('date', '')} • Report ID: {item.get('report_id', '')}{status_text}")
            st.markdown("---")
    else:
        st.info("No projects yet. Go to Projects to start a new intake.")

elif page == "Projects":
    st.markdown('<div class="sy-hero"><div class="sy-hero-copy"><h1>Projects</h1><div class="sy-muted">Create a structured project intake, upload drawings, and generate a professional AI review report.</div></div></div>', unsafe_allow_html=True)
    step = int(st.session_state.get("project_step", 1))
    steps = ["Module", "Details", "Type", "Scope", "Specifics", "Upload", "Report"]
    st.progress(step / len(steps))
    st.caption(" → ".join([f"**{name}**" if i + 1 == step else name for i, name in enumerate(steps)]))
    main_col, side_col = st.columns([1.65, 0.82], gap="large")
    with main_col:
        if step == 1:
            step_header(1, "Choose module", "Select whether this project needs a planning review or a building regulations review.")
            st.session_state["wizard_review_module"] = st.selectbox("Review Module", allowed_review_modules, index=allowed_review_modules.index(st.session_state.get("wizard_review_module", allowed_review_modules[0])))
            st.caption("Downloads are unlocked using credits. Planning PDF = 3 credits. Building Regs PDF = 5 credits. Word export = 1 credit.")
            st.session_state["wizard_review_mode"] = st.selectbox("Report Mode", ["Architect / Professional", "Homeowner Summary"], index=["Architect / Professional", "Homeowner Summary"].index(st.session_state.get("wizard_review_mode", "Architect / Professional")))
            wizard_buttons()
        elif step == 2:
            step_header(2, "Project details", "Add the basic project and site information used in the report cover, council detection and AI context.")
            st.session_state["wizard_project_name"] = st.text_input("Project name", value=st.session_state.get("wizard_project_name", ""), placeholder="e.g. 14 Lyon Road Rear Extension")
            st.session_state["wizard_client_name"] = st.text_input("Client name", value=st.session_state.get("wizard_client_name", ""))
            st.session_state["wizard_project_address"] = st.text_input("Project address", value=st.session_state.get("wizard_project_address", ""))
            st.session_state["wizard_proposal_summary"] = st.text_area("Proposal description", value=st.session_state.get("wizard_proposal_summary", ""), height=120, placeholder="Briefly describe the proposal.")
            wizard_buttons()
        elif step == 3:
            step_header(3, "Project type", "Select the relevant project and property type using clear check boxes.")
            checkbox_grid("Project Type", PROJECT_TYPE_OPTIONS, "wizard_project_types", columns=2)
            if st.session_state.get("wizard_review_module") == "Planning Review":
                single_choice_cards("Property Type", PROPERTY_TYPE_OPTIONS, "wizard_property_type", columns=2)
            else:
                st.session_state["wizard_property_type"] = "Not stated"
                st.info("Property type is mainly used for Planning Review. Building Regulations will focus on technical compliance and uploaded drawings.")
            wizard_buttons()
        elif step == 4:
            step_header(4, "Project scope", "Tick the works included. The AI will cross-check these selections against the uploaded plans and only rely on confirmed drawing information where there is a conflict.")
            checkbox_grid("Scope items", SCOPE_ITEM_OPTIONS, "wizard_scope_items", columns=2)
            if st.session_state.get("wizard_scope_items"):
                st.caption("Selected scope will be passed into the AI context and checked against the uploaded drawings.")
            wizard_buttons()
        elif step == 5:
            step_header(5, "Project specifics", "Enter the key follow-up information. User-entered measurements are reference only; if the uploaded plans show different dimensions, the plans take priority in the report.")
            project_types = st.session_state.get("wizard_project_types", [])
            if st.session_state.get("wizard_review_module") == "Planning Review" and "Ground Floor Rear Extension" in project_types:
                c1, c2 = st.columns(2)
                with c1:
                    st.session_state["wizard_rear_depth"] = st.number_input("Approx. depth from rear wall (metres)", min_value=0.0, max_value=12.0, value=float(st.session_state.get("wizard_rear_depth", 6.0)), step=0.1)
                with c2:
                    st.session_state["wizard_rear_height"] = st.number_input("Approx. maximum height (metres)", min_value=0.0, max_value=6.0, value=float(st.session_state.get("wizard_rear_height", 4.0)), step=0.1)
                st.caption("These dimensions help the initial route check. The uploaded plan measurements take priority if there is any difference.")
            if st.session_state.get("wizard_review_module") == "Planning Review":
                st.session_state["wizard_accuracy_answers"] = get_required_accuracy_answers(project_types)
            else:
                st.session_state["wizard_accuracy_answers"] = {}
                st.info("Building Regulations mode will use the uploaded plans, specification notes and selected scope to focus the compliance review.")
            st.session_state["wizard_review_focus"] = st.text_area(
                "Specific review focus / planning notes",
                value=st.session_state.get("wizard_review_focus", ""),
                height=115,
                placeholder="Example: Focus on PD Class A limits, prior approval risk, Part K stair geometry, Part M accessibility, fire escape strategy, or any specific issue the report should concentrate on."
            )
            wizard_buttons()
        elif step == 6:
            step_header(6, "Upload files", "Upload drawings and supporting information. PDF drawing packs work best for the live review.")
            uploaded = st.file_uploader("Drop files here or click to browse", type=["pdf"], accept_multiple_files=True, key="wizard_pdf_upload")
            if uploaded:
                st.session_state["wizard_uploaded_files"] = uploaded
                for f in uploaded:
                    st.markdown(f'<div class="sy-upload-item"><strong>{f.name}</strong><br><span class="sy-muted">{round(f.size/(1024*1024),2)} MB</span></div>', unsafe_allow_html=True)
                render_pdf_preview(uploaded)
            else:
                st.info("No files attached yet. Upload at least one drawing PDF before generating the report.")
            wizard_buttons()
        elif step == 7:
            step_header(7, "Generate report", "Run the AI review first, then unlock PDF or Word exports using credits.")
            uploaded_files = st.session_state.get("wizard_uploaded_files", [])
            if uploaded_files:
                st.success(f"{len(uploaded_files)} file(s) ready for analysis.")
            else:
                st.warning("No file uploaded yet. Go back to Step 6.")
            if st.button("Analyse Drawing Pack", use_container_width=True):
                run_archlens_analysis(uploaded_files)
            st.markdown("")
            render_report_download_panel(st.session_state.get("active_module", st.session_state.get("wizard_review_module")))
            wizard_buttons()
    with side_col:
        render_intake_panel()

elif page == "Reports":
    st.markdown('<div class="sy-hero"><div class="sy-hero-copy"><h1>Reports</h1><div class="sy-muted">Your generated report library. Download previous Planning and Building Regulations reports again without showing the full report output.</div></div></div>', unsafe_allow_html=True)
    saved_projects = st.session_state.get("saved_projects", [])
    if saved_projects:
        for item in saved_projects:
            st.markdown('<div class="sy-report-card">', unsafe_allow_html=True)
            c1, c2, c3 = st.columns([1.4, 0.9, 0.9])
            with c1:
                st.markdown(f"**{item.get('project_address', 'Not provided')}**")
                st.caption(f"Report ID: {item.get('report_id', '')} • {item.get('filename', '')}")
            with c2:
                st.write(item.get("module", ""))
                conf = item.get("ai_confidence") or {}
                if conf.get("label"):
                    st.caption(f"Status: {conf.get('label')}")
                st.caption(f"Council: {item.get('local_authority', 'Not detected')}")
            with c3:
                st.write(item.get("date", ""))
                st.caption(item.get("plan", ""))
            d1, d2 = st.columns(2)
            pdf_bytes = item.get("pdf_bytes")
            word_bytes = item.get("word_bytes")
            report_id_item = item.get("report_id", "")
            module_item = item.get("module", "Planning Review")
            with d1:
                if pdf_bytes and item.get("pdf_unlocked"):
                    st.download_button("Download PDF", pdf_bytes, file_name=f"{report_id_item or 'report'}_ArchLens_Report.pdf", mime="application/pdf", use_container_width=True, key=f"pdf_{report_id_item}")
                elif pdf_bytes:
                    cost = get_export_credit_cost(module_item, "pdf")
                    if st.button(f"Unlock PDF — {cost} credits", use_container_width=True, key=f"unlock_pdf_{report_id_item}"):
                        ok, msg = unlock_report_export(report_id_item, module_item, "pdf")
                        if ok:
                            st.success(msg)
                            st.rerun()
                        else:
                            st.error(msg)
                else:
                    st.caption("PDF download available for reports generated after this update.")
            with d2:
                if word_bytes and item.get("word_unlocked"):
                    st.download_button("Download Word", word_bytes, file_name=f"{report_id_item or 'report'}_ArchLens_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key=f"docx_{report_id_item}")
                elif word_bytes:
                    cost = get_export_credit_cost(module_item, "word")
                    if st.button(f"Unlock Word — {cost} credit", use_container_width=True, key=f"unlock_word_{report_id_item}"):
                        ok, msg = unlock_report_export(report_id_item, module_item, "word")
                        if ok:
                            st.success(msg)
                            st.rerun()
                        else:
                            st.error(msg)
                else:
                    st.caption("Word download available for reports generated after this update.")
            st.markdown('</div>', unsafe_allow_html=True)
    else:
        st.info("No reports generated yet. Go to Projects and run your first review.")

elif page == "Settings":
    st.markdown('<div class="sy-hero"><div class="sy-hero-copy"><h1>Settings</h1><div class="sy-muted">Control your account, branding and app appearance.</div></div></div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("### Account")
        st.write(f"Current plan: **{PLAN_LABELS.get(current_plan, 'Solo')}**")
        st.write(f"User: **{current_user_name or 'Not shown'}**")
        st.write(f"Credits available: **{get_credit_balance()}**")
        st.caption("This version uses session-based credits for testing. Live credits should be moved to Supabase/Stripe so balances persist.")
        st.markdown("### Appearance")
        selected_theme = st.radio("App theme", ["Dark", "Light"], index=["Dark", "Light"].index(st.session_state.get("app_theme", "Dark")), horizontal=True)
        if selected_theme != st.session_state.get("app_theme", "Dark"):
            st.session_state["app_theme"] = selected_theme
            st.rerun()
        st.caption("Light mode uses dark text on light surfaces. Dark mode uses light text on dark surfaces for readability.")
    with c2:
        st.markdown("### Branding")
        logo_bytes = get_brand_logo_bytes_for_ui()
        if logo_bytes:
            st.image(logo_bytes, width=220)
            st.success("SY Design Studio logo is loaded for branded PDF exports.")
        else:
            st.warning("Logo file not found. Add assets/sy_design_studio_logo.png to your project.")
    st.markdown("---")
    render_buy_credits_panel()

    st.markdown("---")
    st.markdown("### Credit Transactions")
    transactions = st.session_state.get("credit_transactions", []) or []
    if transactions:
        for tx in transactions[:8]:
            sign = "+" if int(tx.get("amount", 0)) > 0 else ""
            st.caption(f"{tx.get('date')} • {sign}{tx.get('amount')} credits • {tx.get('reason')} • Balance: {tx.get('balance_after')}")
    else:
        st.caption("No credit transactions yet.")

    if st.button("Clear current project/report"):
        # Do not reset credit balance, credit transactions, unlocked reports or report library.
        # Credits are money-related and must persist in the user session.
        preserve_keys = {"credit_balance", "credit_transactions", "unlocked_reports", "saved_projects", "report_library"}
        for key, value in DEFAULT_STATE.items():
            if key not in preserve_keys:
                st.session_state[key] = value
        for key in list(st.session_state.keys()):
            if key.startswith("wizard_"):
                st.session_state[key] = WIZARD_DEFAULTS.get(key, "")
        st.session_state["project_step"] = 1
        st.success("Current project cleared. Credits and previously generated report library were kept.")
